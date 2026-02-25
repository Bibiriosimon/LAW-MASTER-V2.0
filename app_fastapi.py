"""
LAW MASTER v1.0 - FastAPI local UI (ChatGPT-like)
Run:
  python -m uvicorn app_fastapi:app --host 127.0.0.1 --port 7860 --reload
"""
# -*- coding: utf-8 -*-

import json
import os
import time
import logging
import re
import unicodedata
from concurrent.futures import ThreadPoolExecutor
import uuid
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import jieba
import numpy as np
import requests
from bs4 import BeautifulSoup
from urllib.parse import urlparse, urljoin
from readability import Document
from ddgs import DDGS
from threading import Lock
from concurrent.futures import ThreadPoolExecutor, as_completed
from charset_normalizer import from_bytes
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from dotenv import load_dotenv
from fastapi import FastAPI, File, UploadFile, Body, Form
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# -----------------------------
# Paths & logging
# -----------------------------
PROJECT_ROOT = Path(__file__).resolve().parent
CORPUS_PATH = PROJECT_ROOT / "corpus_enriched_fast.jsonl"
EXT_KB_PATH = PROJECT_ROOT / "corpus_extension.jsonl"
CASE_KB_PATH = PROJECT_ROOT / "case_kb.jsonl"
if not CASE_KB_PATH.exists():
    alt_case = PROJECT_ROOT.parent / "case_kb.jsonl"
    if alt_case.exists():
        CASE_KB_PATH = alt_case
    else:
        alt_case2 = PROJECT_ROOT.parent.parent / "case_kb.jsonl"
        if alt_case2.exists():
            CASE_KB_PATH = alt_case2
CACHE_DIR = PROJECT_ROOT / "cache"
LOG_DIR = PROJECT_ROOT / "logs"
WEB_DIR = PROJECT_ROOT / "web"
EXPORT_DIR = PROJECT_ROOT / "exports"
ASSET_DIR = CACHE_DIR / "assets"
MEMORY_DIR = CACHE_DIR / "memory"
STATE_DIR = CACHE_DIR / "memory_state"

CACHE_DIR.mkdir(parents=True, exist_ok=True)
LOG_DIR.mkdir(parents=True, exist_ok=True)
EXPORT_DIR.mkdir(parents=True, exist_ok=True)
ASSET_DIR.mkdir(parents=True, exist_ok=True)
MEMORY_DIR.mkdir(parents=True, exist_ok=True)
STATE_DIR.mkdir(parents=True, exist_ok=True)

# search diagnostics
SEARCH_DIAG = {
    "engine": "ddgs",
    "backends": "",
    "region": "",
    "queries": [],
    "total_results": 0,
    "content_hits": 0,
    "ts": 0,
}
SEARCH_DIAG_LOCK = Lock()

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    handlers=[
        logging.FileHandler(LOG_DIR / "app_fastapi.log", encoding="utf-8"),
        logging.StreamHandler(),
    ],
)
logger = logging.getLogger("lawmaster_fastapi")

# -----------------------------
# Config
# -----------------------------
@dataclass
class AppConfig:
    name: str = "LAW MASTER v1.0"
    base_url: str = "https://api.deepseek.com/v1"
    model: str = "deepseek-chat"
    top_k: int = 5
    keyword_weight: float = 0.7
    vector_weight: float = 0.3
    max_context_chars: int = 4000
    max_file_chars: int = 2000
    rewrite_temperature: float = 0.2
    answer_temperature: float = 0.6
    answer_max_tokens: int = 3000


CONFIG = AppConfig()

# -----------------------------
# Env defaults for OCR stability
# -----------------------------
os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
os.environ.setdefault("FLAGS_enable_pir_api", "0")
os.environ.setdefault("FLAGS_enable_pir_in_executor", "0")
os.environ.setdefault("FLAGS_use_mkldnn", "0")
os.environ.setdefault("FLAGS_enable_mkldnn", "0")
os.environ.setdefault("FLAGS_enable_onednn", "0")

# -----------------------------
# Document templates
# -----------------------------

DOCUMENT_TEMPLATES = {
    "申诉书": [
        "标题：申诉书",
        "致：{受理机构/部门}",
        "申诉人信息：姓名、身份证号、联系方式、地址",
        "案由/事项：{简述争议事项}",
        "事实经过：{时间线、关键事实}",
        "申诉理由：{依据、程序问题、事实不清等}",
        "诉求：{请求撤销/复核/改正}",
        "证据清单：{证据名称、编号}",
        "落款：{签名/日期}",
    ],
    "投诉信": [
        "标题：投诉信",
        "致：{投诉受理部门}",
        "投诉人信息：姓名、联系方式、地址",
        "被投诉对象：{机构/人员}",
        "投诉事项：{核心问题}",
        "事实经过：{时间、地点、具体行为}",
        "诉求：{整改/赔偿/解释}",
        "附件清单：{证据材料}",
        "落款：{签名/日期}",
    ],
    "律师函": [
        "标题：律师函",
        "函件编号：{编号}",
        "致：{收函对象}",
        "委托人信息：{姓名/公司}",
        "事实与背景：{简述事实}",
        "法律依据：{法律条文/合同条款}",
        "要求与期限：{具体要求、期限}",
        "后果提示：{不履行的法律后果}",
        "律师信息：{律所、联系方式}",
        "落款：{律师签名/日期}",
    ],
    "起诉状": [
        "标题：民事起诉状",
        "原告信息：{姓名/单位、地址、联系方式}",
        "被告信息：{姓名/单位、地址、联系方式}",
        "诉讼请求：{请求事项}",
        "事实与理由：{事实经过、证据、法律依据}",
        "证据目录：{证据名称、编号}",
        "附项：{管辖法院、诉讼费用}",
        "落款：{签名/日期}",
    ],
    "通用文书": [
        "标题：{文书标题}",
        "收件人：{机构/部门}",
        "当事人信息：{姓名/联系方式}",
        "事实描述：{时间线/关键事实}",
        "诉求/目的：{期望结果}",
        "证据：{证据清单}",
        "落款：{签名/日期}",
    ],
}

# -----------------------------
# Utilities
# -----------------------------


def safe_json_loads(text: str) -> Optional[Dict[str, Any]]:
    try:
        return json.loads(text)
    except Exception:
        return None


def extract_json_block(text: str) -> Optional[Dict[str, Any]]:
    if not text:
        return None
    candidate = extract_first_json_object(text) or ""
    if not candidate:
        return None
    return safe_json_loads(candidate)


def jieba_tokenize(text: str) -> List[str]:
    return [t.strip() for t in jieba.lcut(text) if t.strip()]


def fallback_keywords(text: str, k: int = 6) -> List[str]:
    tokens = [t for t in jieba.lcut(text) if t.strip() and len(t) > 1]
    counts: Dict[str, int] = {}
    for t in tokens:
        counts[t] = counts.get(t, 0) + 1
    ranked = sorted(counts.items(), key=lambda x: x[1], reverse=True)
    return [w for w, _ in ranked[:k]]


def truncate(text: str, limit: int) -> str:
    if text is None:
        return ""
    return text if len(text) <= limit else text[:limit] + "..."


def is_image_ext(ext: str) -> bool:
    return ext.lower() in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]


def save_upload_asset(filename: str, content: bytes) -> str:
    asset_id = uuid.uuid4().hex
    suffix = Path(filename).suffix.lower() or ".bin"
    path = ASSET_DIR / f"{asset_id}{suffix}"
    path.write_bytes(content)
    return asset_id


def find_asset_path(asset_id: str) -> Optional[Path]:
    if not asset_id:
        return None
    matches = list(ASSET_DIR.glob(f"{asset_id}.*"))
    if matches:
        return matches[0]
    return None


def memory_path(session_id: str) -> Path:
    safe = re.sub(r"[^a-zA-Z0-9_-]+", "_", session_id or "default")
    return MEMORY_DIR / f"{safe}.jsonl"


def state_path(session_id: str) -> Path:
    safe = re.sub(r"[^a-zA-Z0-9_-]+", "_", session_id or "default")
    return STATE_DIR / f"{safe}.json"


def load_state(session_id: str) -> Dict[str, Any]:
    path = state_path(session_id)
    if not path.exists():
        return {
            "pending_clarify": False,
            "pending_for": "",
            "pending_questions": [],
            "skip_clarify": False,
        }
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return {
            "pending_clarify": False,
            "pending_for": "",
            "pending_questions": [],
            "skip_clarify": False,
        }


def save_state(session_id: str, state: Dict[str, Any]) -> None:
    try:
        state_path(session_id).write_text(json.dumps(state, ensure_ascii=False), encoding="utf-8")
    except Exception:
        pass


def load_memory(session_id: str, limit: int = 8) -> List[Dict[str, Any]]:
    path = memory_path(session_id)
    if not path.exists():
        return []
    items: List[Dict[str, Any]] = []
    try:
        with path.open("r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line:
                    continue
                try:
                    items.append(json.loads(line))
                except Exception:
                    continue
    except Exception:
        return []
    return items[-limit:]


def append_memory(session_id: str, entry: Dict[str, Any]) -> None:
    path = memory_path(session_id)
    try:
        with path.open("a", encoding="utf-8") as f:
            f.write(json.dumps(entry, ensure_ascii=False) + "\n")
    except Exception:
        pass


def build_memory_context(session_id: str) -> str:
    mem = load_memory(session_id)
    if not mem:
        return ""
    lines = []
    for m in mem:
        role = m.get("role", "")
        content = m.get("content", "")
        if content:
            lines.append(f"{role}:{content}")
    return "\n".join(lines)


def extract_sparse_memory(user_text: str, assistant_text: str, doc_titles: Optional[List[str]] = None) -> List[Dict[str, Any]]:
    items = []
    user_text = clean_text(user_text)[:240]
    assistant_text = clean_text(assistant_text)[:240]
    if user_text:
        items.append({"role": "user", "content": user_text, "ts": int(time.time())})
    if assistant_text:
        items.append({"role": "assistant", "content": assistant_text, "ts": int(time.time())})
    if doc_titles:
        items.append({"role": "assistant", "content": f"已生成文书: {', '.join(doc_titles)}", "ts": int(time.time())})
    return items


def annotate_plan_detail(
    plan_detail: Dict[str, Any],
    rag_keywords: Optional[List[str]] = None,
    search_results: Optional[List[Dict[str, Any]]] = None,
    documents: Optional[List[Dict[str, Any]]] = None,
    action: str = "",
    kb_brief: Optional[List[Dict[str, Any]]] = None,
    web_brief: Optional[List[Dict[str, Any]]] = None,
) -> Dict[str, Any]:
    if not plan_detail or not isinstance(plan_detail, dict):
        return plan_detail
    steps = plan_detail.get("steps") or []
    if not steps:
        return plan_detail
    if len(steps) >= 1:
        kw = ", ".join(rag_keywords or [])[:60]
        steps[0]["result"] = f"关键词：{kw or '无'}"
        if kb_brief:
            steps[0]["inspect"] = {"kb": kb_brief}
    if len(steps) >= 2:
        sr = search_results or []
        if sr:
            top = []
            for r in sr[:2]:
                title = r.get("title") or "网页"
                host = urlparse(r.get("url", "")).netloc
                top.append(f"{title} ({host})")
            steps[1]["result"] = f"网页命中：{len(sr)} | 摘要：{'；'.join(top)}"
        else:
            steps[1]["result"] = "网页命中：0"
        if web_brief:
            steps[1]["inspect"] = {"web": web_brief}
    if len(steps) >= 3:
        docs = documents or []
        if action == "generate_document" or docs:
            steps[2]["result"] = f"文书生成：{len(docs)} 份"
        else:
            steps[2]["result"] = "完成答复"
    plan_detail["steps"] = steps
    return plan_detail


def clean_text(text: str) -> str:
    if not text:
        return ""
    # remove replacement chars and control chars
    text = text.replace("\ufffd", "")
    text = "".join(ch for ch in text if ch >= " " or ch in "\n\t")
    # normalize whitespace
    text = "\n".join([line.strip() for line in text.splitlines() if line.strip()])
    return text.strip()


def sanitize_answer_text(text: str) -> str:
    if not text:
        return ""
    t = clean_text(text)
    # remove asterisks
    t = t.replace("*", "")
    return t


def sanitize_session_title(text: str) -> str:
    if not text:
        return ""
    t = sanitize_answer_text(text).strip()
    t = t.splitlines()[0] if t else ""
    t = re.sub(r"^\s*[#>*\-]+\s*", "", t)
    t = re.sub(r"^\s*(?:\d+[\.\)\-、]\s*)+", "", t)
    t = re.sub(r"^\s*[一二三四五六七八九十]+\s*[、\.．]\s*", "", t)
    t = re.sub(r"^(会话标题|标题|建议标题|结论|摘要|总结|分析)\s*[:：\-]?\s*", "", t, flags=re.I)
    t = t.strip().strip("`\"'“”[]【】()（）")
    t = re.sub(r"\s+", "", t)
    t = re.sub(r"[。！？!?；;，,：:]+$", "", t)
    t = re.sub(r"(问题|咨询|结论|摘要)$", "", t)
    if len(t) < 2:
        return ""
    return t[:16]


def fallback_session_title_from_query(query: str) -> str:
    q = (query or "").strip()
    if not q:
        return "新聊天"
    q = re.sub(r"\s+", "", q)
    q = re.sub(r"^(请问|咨询|我想问|我想咨询|我想了解|我今天|我现在|我最近|我|本人)", "", q)
    q = re.sub(r"(怎么办|怎么处理|是否合法|违法吗|可以吗|能否|吗|呢)[？?。！!]*$", "", q)
    q = re.sub(r"[，,。！？!?；;:：、\"'“”`]", "", q)
    q = q.strip()
    if len(q) < 2:
        return "新聊天"
    if q.endswith("问题"):
        q = q[:-2]
    return q[:12] or "新聊天"


def looks_mojibake(text: str) -> bool:
    if not text:
        return False
    total = len(text)
    if total < 80:
        return False
    latin1 = sum(1 for ch in text if "\u00c0" <= ch <= "\u00ff")
    cjk = sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")
    return latin1 / max(total, 1) > 0.15 and cjk / max(total, 1) < 0.1


def _cjk_score(text: str) -> int:
    return sum(1 for ch in text if "\u4e00" <= ch <= "\u9fff")


def fix_mojibake(text: str) -> str:
    if not text or not looks_mojibake(text):
        return text
    candidates = []
    for enc in ("latin1", "cp1252"):
        try:
            candidates.append(text.encode(enc).decode("utf-8"))
        except Exception:
            pass
        try:
            candidates.append(text.encode(enc).decode("gbk"))
        except Exception:
            pass
    if not candidates:
        return text
    best = max(candidates, key=_cjk_score)
    return best if _cjk_score(best) >= _cjk_score(text) else text


def clean_ocr_text(text: str) -> str:
    if not text:
        return ""
    t = unicodedata.normalize("NFKC", text)
    t = clean_text(t)
    # remove common OCR noise patterns
    t = re.sub(r"[|¦·•]{3,}", "", t)
    # collapse extra spaces
    t = re.sub(r"[^\S\n]+", " ", t)
    # remove spaces between CJK characters
    for _ in range(2):
        t = re.sub(r"([\u4e00-\u9fff])\s+([\u4e00-\u9fff])", r"\1\2", t)
    lines = []
    for line in t.splitlines():
        line = line.strip()
        if not line:
            continue
        # drop lines that are almost all punctuation
        useful = sum(1 for ch in line if ch.isalnum() or ("\u4e00" <= ch <= "\u9fff"))
        if useful == 0 and len(line) < 6:
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def generate_case_keywords(title: str, content: str, top_n: int = 10) -> List[str]:
    text = f"{title}\n{content}"
    tokens = [t.strip() for t in jieba.lcut(text) if t.strip()]
    counts: Dict[str, int] = {}
    for t in tokens:
        if len(t) < 2:
            continue
        if re.fullmatch(r"\d+", t):
            continue
        weight = 3 if title and t in title else 1
        counts[t] = counts.get(t, 0) + weight
    ranked = sorted(counts.items(), key=lambda x: x[1], reverse=True)
    return [w for w, _ in ranked[:top_n]]


def build_case_summary(text: str, max_chars: int = 360) -> str:
    if not text:
        return ""
    parts = re.split(r"[。！？!?]\s*", text)
    summary = "。".join([p for p in parts if p][:3]).strip()
    if not summary:
        summary = text[:max_chars]
    if len(summary) > max_chars:
        summary = summary[:max_chars] + "..."
    return summary


def is_gibberish(text: str) -> bool:
    if not text:
        return True
    if looks_mojibake(text):
        return True
    if len(text) < 120:
        return False
    total = len(text)
    useful = 0
    for ch in text:
        if ch.isalnum() or ("\u4e00" <= ch <= "\u9fff"):
            useful += 1
    return useful / max(total, 1) < 0.28


def extract_first_json_object(text: str) -> Optional[str]:
    if not text:
        return None
    start = None
    depth = 0
    in_str = False
    escape = False
    for i, ch in enumerate(text):
        if start is None:
            if ch == "{":
                start = i
                depth = 1
            continue
        if in_str:
            if escape:
                escape = False
            elif ch == "\\":
                escape = True
            elif ch == '"':
                in_str = False
            continue
        if ch == '"':
            in_str = True
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0 and start is not None:
                return text[start : i + 1]
    return None


def decode_html_bytes(content: bytes) -> str:
    detected = from_bytes(content).best()
    if detected:
        text = str(detected)
        if text and not looks_mojibake(text):
            return text
    # fallback encodings for CN sites
    for enc in ("utf-8", "gb18030", "gbk", "gb2312", "big5"):
        try:
            text = content.decode(enc, errors="ignore")
            if text and not looks_mojibake(text):
                return text
        except Exception:
            continue
    return content.decode("utf-8", errors="ignore")


def normalize_str_list(val: Any) -> List[str]:
    if isinstance(val, list):
        return [str(x).strip() for x in val if str(x).strip()]
    if isinstance(val, str):
        parts = [p.strip() for p in re.split(r"[；;。\n]+", val) if p.strip()]
        return parts
    return []


def normalize_subtasks(val: Any) -> List[List[str]]:
    if isinstance(val, list):
        out: List[List[str]] = []
        for item in val:
            if isinstance(item, list):
                out.append([str(x).strip() for x in item if str(x).strip()])
            elif isinstance(item, str):
                out.append([p.strip() for p in re.split(r"[；;。\n]+", item) if p.strip()])
            else:
                out.append([])
        return out
    if isinstance(val, str):
        parts = [p.strip() for p in re.split(r"[；;。\n]+", val) if p.strip()]
        return [parts] if parts else []
    return []


def normalize_model_json(text: str) -> str:
    if not text:
        return ""
    t = text.strip()
    # strip code fences
    if t.startswith("```"):
        t = re.sub(r"^```[a-zA-Z]*\n", "", t)
        t = t.rstrip("`").strip()
    # strip leading 'json' token
    if t.lower().startswith("json"):
        t = t[4:].lstrip()
    return t


def extract_answer_field(text: str) -> str:
    if not text:
        return ""
    # try structured parse first
    candidate = extract_first_json_object(text) or text
    data = safe_json_loads(candidate)
    if isinstance(data, dict) and data.get("answer"):
        return str(data.get("answer"))
    # strict quoted answer
    m = re.search(r'"answer"\s*:\s*"((?:\\.|[^"\\])*)"', text, re.S)
    if m:
        raw = m.group(1)
        try:
            return json.loads('"' + raw + '"')
        except Exception:
            return raw
    # loose answer until next field
    m = re.search(r'"answer"\s*:\s*(.+?)(?:"\s*,\s*"(?:kb_entries|documents)"|\n\}|\}$)', text, re.S)
    if m:
        raw = m.group(1).strip().strip(",")
        raw = raw.strip('"')
        return raw
    return ""


# -----------------------------
# Corpus & index
# -----------------------------


def load_corpus(path: Path) -> List[Dict[str, Any]]:
    docs: List[Dict[str, Any]] = []
    if not path.exists():
        return docs
    raw = path.read_bytes()
    try:
        text = raw.decode("utf-8")
    except Exception:
        match = from_bytes(raw).best()
        text = match.output if match else raw.decode("utf-8", errors="ignore")
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            doc = json.loads(line)
        except Exception:
            logger.warning("Skipped invalid JSON line")
            continue
        if isinstance(doc, dict):
            article = fix_mojibake(clean_text(doc.get("article_number", "")))
            content = fix_mojibake(clean_text(doc.get("content", "")))
            source = fix_mojibake(clean_text(doc.get("source", "")))
            doc["article_number"] = article
            doc["content"] = content
            doc["source"] = source
            doc["doc_type"] = "law"
        docs.append(doc)
    return docs


def load_extension_corpus(path: Path) -> List[Dict[str, Any]]:
    if not path.exists():
        return []
    return load_corpus(path)


def load_case_corpus(path: Path) -> List[Dict[str, Any]]:
    docs: List[Dict[str, Any]] = []
    if not path.exists():
        return docs
    raw = path.read_bytes()
    try:
        text = raw.decode("utf-8")
    except Exception:
        match = from_bytes(raw).best()
        text = match.output if match else raw.decode("utf-8", errors="ignore")
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        try:
            data = json.loads(line)
        except Exception:
            continue
        if not isinstance(data, dict):
            continue
        title = fix_mojibake(clean_text(data.get("title") or data.get("case_title") or ""))
        content = fix_mojibake(clean_text(data.get("text") or data.get("content") or ""))
        source = fix_mojibake(clean_text(data.get("source") or data.get("source_file") or ""))
        if not content:
            continue
        keywords = data.get("keywords") or generate_case_keywords(title, content)
        summary = build_case_summary(content)
        doc = {
            "id": data.get("id") or f"case_{len(docs)}",
            "title": title or "案例",
            "content": content,
            "source": source,
            "page_start": data.get("page_start"),
            "page_end": data.get("page_end"),
            "keywords": keywords,
            "summary": summary,
            "doc_type": "case",
        }
        docs.append(doc)
    return docs


def build_doc_text(doc: Dict[str, Any]) -> str:
    doc_type = doc.get("doc_type") or "law"
    if doc_type == "case":
        parts = [
            doc.get("title", ""),
            doc.get("title", ""),
            doc.get("content", ""),
            " ".join(doc.get("keywords", []) or []),
            doc.get("source", ""),
        ]
    else:
        parts = [
            doc.get("article_number", ""),
            doc.get("content", ""),
            " ".join(doc.get("keywords", []) or []),
            " ".join(doc.get("hypothetical_questions", []) or []),
            doc.get("source", ""),
        ]
    return "\n".join([p for p in parts if p])


def get_cache_key(path: Path) -> str:
    stat = path.stat()
    return f"{stat.st_size}_{int(stat.st_mtime)}"


def get_ext_cache_key() -> str:
    if EXT_KB_PATH.exists():
        stat = EXT_KB_PATH.stat()
        return f"_{stat.st_size}_{int(stat.st_mtime)}"
    return "_noext"


def load_or_build_index(docs: List[Dict[str, Any]], corpus_path: Path, extra_key: str = ""):
    cache_key = get_cache_key(corpus_path)
    cache_path = CACHE_DIR / f"tfidf_{cache_key}{extra_key}.joblib"

    if cache_path.exists():
        try:
            import joblib

            data = joblib.load(cache_path)
            return data["vectorizer"], data["matrix"], data["doc_texts"]
        except Exception:
            logger.warning("Cache load failed; rebuilding index")

    doc_texts = [build_doc_text(d) for d in docs]
    vectorizer = TfidfVectorizer(
        tokenizer=jieba_tokenize,
        token_pattern=None,
        max_features=200000,
        lowercase=False,
    )
    matrix = vectorizer.fit_transform(doc_texts)

    try:
        import joblib

        joblib.dump({"vectorizer": vectorizer, "matrix": matrix, "doc_texts": doc_texts}, cache_path)
    except Exception:
        logger.warning("Cache save failed")

    return vectorizer, matrix, doc_texts


def rebuild_index():
    global VECTORIZER, MATRIX, _DOC_TEXTS
    VECTORIZER, MATRIX, _DOC_TEXTS = load_or_build_index(
        DOCS, CORPUS_PATH, extra_key="_with_ext" + get_ext_cache_key()
    )


def rebuild_case_index():
    global CASE_VECTORIZER, CASE_MATRIX, _CASE_DOC_TEXTS
    if not CASE_DOCS:
        CASE_VECTORIZER, CASE_MATRIX, _CASE_DOC_TEXTS = None, None, []
        return
    CASE_VECTORIZER, CASE_MATRIX, _CASE_DOC_TEXTS = load_or_build_index(
        CASE_DOCS, CASE_KB_PATH, extra_key="_case"
    )


# -----------------------------
# DeepSeek API
# -----------------------------


def call_deepseek(
    api_key: str,
    base_url: str,
    model: str,
    messages: List[Dict[str, str]],
    temperature: float,
    max_tokens: int,
) -> str:
    url = base_url.rstrip("/") + "/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "Connection": "close",
    }
    payload = {
        "model": model,
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens,
    }
    last_err = None
    for attempt in range(3):
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=(10, 120))
            resp.raise_for_status()
            try:
                data = resp.json()
            except Exception:
                data = json.loads(resp.content.decode("utf-8", errors="replace"))
            return data["choices"][0]["message"]["content"]
        except requests.HTTPError as e:
            raise RuntimeError(f"DeepSeek API 错误: {resp.status_code} {resp.text}") from e
        except requests.RequestException as e:
            last_err = e
            time.sleep(0.8 + attempt * 0.8)
    raise RuntimeError(f"DeepSeek API 请求失败: {last_err}") from last_err


def call_reasoner(
    api_key: str,
    base_url: str,
    model: str,
    messages: List[Dict[str, str]],
    temperature: float = 0.2,
    max_tokens: int = 900,
) -> str:
    return call_deepseek(
        api_key=api_key,
        base_url=base_url,
        model=model,
        messages=messages,
        temperature=temperature,
        max_tokens=max_tokens,
    )


def plan_self_awareness(
    api_key: str,
    base_url: str,
    model: str,
    history: List[Dict[str, str]],
    memory_context: str = "",
) -> Dict[str, Any]:
    system_prompt = (
        "你是规划型智能体(Deep Research)。只输出JSON。"
        "目标：判断是否需要澄清、制定探索计划、产出搜索关键词/查询、决定是否需要更新知识库。"
        "当用户请求具有多目标/不明确选项时，优先给出1-2个澄清问题。"
        "当用户明确要求生成文书（申诉书/投诉信/律师函/起诉状等）时，可将 action 设置为 generate_document，"
        "并把 document_type 填为对应类型，同时 search_needed=false。"
    )
    history_text = "\n".join([f"{m.get('role')}: {m.get('content')}" for m in history[-6:]])
    user_prompt = (
        "对话历史如下：\n"
        f"{history_text}\n\n"
        f"历史要点（稀疏记忆）：\n{memory_context or '(无)'}\n\n"
        "请输出JSON，格式如下：\n"
        "{\n"
        '  "need_more_info": false,\n'
        '  "clarifying_questions": ["问题1","问题2"],\n'
        '  "search_needed": true,\n'
        '  "search_queries": ["查询1","查询2"],\n'
        '  "action": "rag|web_search|generate_document|direct_answer",\n'
        '  "document_type": "申诉书/投诉信/律师函/起诉状/通用文书",\n'
        '  "goal_analysis": "目标分析与范围界定",\n'
        '  "plan": ["步骤1","步骤2","步骤3"],\n'
        '  "step_details": ["细节1","细节2","细节3"],\n'
        '  "subtasks": [["子任务1","子任务2"],["子任务1"],[]],\n'
        '  "should_update_kb": true\n'
        "}\n"
        "要求：简洁、可执行。"
    )
    text = call_reasoner(
        api_key=api_key,
        base_url=base_url,
        model=model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
    )
    data = extract_json_block(text) or safe_json_loads(text) or {}
    if not isinstance(data, dict):
        data = {}
    return {
        "need_more_info": bool(data.get("need_more_info", False)),
        "clarifying_questions": data.get("clarifying_questions", []) or [],
        "search_needed": bool(data.get("search_needed", False)),
        "search_queries": data.get("search_queries", []) or [],
        "action": data.get("action", "") or "",
        "document_type": data.get("document_type", "") or "",
        "goal_analysis": data.get("goal_analysis", "") or "",
        "plan": data.get("plan", []) or [],
        "step_details": data.get("step_details", []) or [],
        "subtasks": data.get("subtasks", []) or [],
        "should_update_kb": bool(data.get("should_update_kb", False)),
    }


def refine_search_queries(api_key: str, base_url: str, model: str, user_query: str, prev_queries: List[str]) -> List[str]:
    system_prompt = "你是搜索关键词优化助手，只输出JSON。"
    user_prompt = (
        f"用户问题：{user_query}\n"
        f"已有查询：{prev_queries}\n"
        "请输出JSON：\n"
        "{ \"queries\": [\"新查询1\",\"新查询2\",\"新查询3\"] }\n"
        "要求：更具体，包含权威来源关键词。"
    )
    text = call_deepseek(
        api_key=api_key,
        base_url=base_url,
        model=model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=0.2,
        max_tokens=300,
    )
    data = extract_json_block(text) or safe_json_loads(text) or {}
    if not isinstance(data, dict):
        data = {}
    queries = data.get("queries", []) or []
    return [q for q in queries if q.strip()]


def heuristic_queries(user_query: str) -> List[str]:
    tokens = [t for t in jieba.lcut(user_query) if len(t.strip()) > 1]
    base = " ".join(tokens[:10]) if tokens else user_query
    return [
        base,
        base + " 法律规定",
        base + " 维权 处理 方式",
        base + " 典型案例",
        base + " 官方 规则",
    ]


LEGAL_HINTS = (
    "法", "法律", "法规", "条例", "规定", "条款", "条文", "法条", "判决", "裁判", "法院",
    "起诉", "立案", "诉讼", "诉讼时效", "仲裁", "合同", "违约", "侵权", "赔偿", "欺诈",
    "消费者", "维权", "投诉", "申诉", "律师函", "起诉状", "行政", "刑事", "民事",
    "平台规则", "售后", "退货", "退款", "赔付", "证据", "举证", "损失",
    "商家", "平台", "订单", "网购", "电商", "换货", "质量", "保修", "发票", "差评",
    "京东", "淘宝", "拼多多", "抖音", "快手",
)

NON_LEGAL_HINTS = (
    "礼物", "送礼", "情人节", "生日", "纪念日", "惊喜",
    "旅游", "旅行", "机票", "酒店", "民宿",
    "美食", "餐厅", "菜谱", "减肥", "健身",
    "电影", "电视剧", "综艺", "小说", "游戏",
    "恋爱", "表白", "相亲", "情感",
    "穿搭", "护肤", "化妆", "发型",
    "推荐", "选什么", "买什么", "怎么选",
)


def is_legal_query(text: str) -> bool:
    if not text:
        return False
    for kw in LEGAL_HINTS:
        if kw in text:
            return True
    return False


def is_non_legal_query(text: str) -> bool:
    if not text:
        return False
    for kw in NON_LEGAL_HINTS:
        if kw in text:
            return True
    return False


ALLOWED_TLDS = (".cn",)
ALLOWED_DOMAINS = (".gov.cn", ".edu.cn", ".csc.edu.cn", ".csu.edu.cn")


def is_cn_site(url: str) -> bool:
    try:
        host = urlparse(url).netloc.lower()
        if any(host.endswith(d) for d in ALLOWED_DOMAINS):
            return True
        return host.endswith(ALLOWED_TLDS)
    except Exception:
        return False


def filter_by_region(url: str, region: str) -> bool:
    if region == "global":
        return True
    if region == "cn_only":
        return is_cn_site(url)
    if region == "cn_priority":
        # allow all, but caller should prefer cn sites
        return True
    return True


def fetch_page_text(url: str, max_chars: int = 4000) -> str:
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept-Language": "zh-CN,zh;q=0.9",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        }
        sess = requests.Session()
        # warm up cookies for some strict sites
        if "csc.edu.cn" in url:
            try:
                sess.get("https://www.csc.edu.cn", headers=headers, timeout=8)
            except Exception:
                pass
        resp = sess.get(url, headers=headers, timeout=12)
        if resp.status_code == 412:
            # retry with referer & accept-language
            headers.update(
                {
                    "Referer": "https://www.csc.edu.cn/",
                    "Accept-Language": "zh-CN,zh;q=0.9",
                    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                }
            )
            resp = sess.get(url, headers=headers, timeout=12)
        resp.raise_for_status()
        # skip non-html
        if "text/html" not in resp.headers.get("Content-Type", ""):
            return ""
        html = decode_html_bytes(resp.content)
        if "bing.com/search" in resp.url or "Bing" in html[:5000]:
            return ""
        doc = Document(html)
        summary_html = doc.summary()
        soup = BeautifulSoup(summary_html, "lxml")
        text = soup.get_text("\n", strip=True)
        if not text:
            soup = BeautifulSoup(html, "lxml")
            text = soup.get_text("\n", strip=True)
        text = clean_text(text)
        if is_gibberish(text):
            return ""
        return text[:max_chars]
    except Exception as e:
        logger.warning("Fetch page failed: %s %s", url, e)
        return ""


def extract_internal_links(base_url: str, html: str, limit: int = 4) -> List[str]:
    links = []
    try:
        base_host = urlparse(base_url).netloc
        soup = BeautifulSoup(html, "lxml")
        for a in soup.select("a[href]"):
            href = a.get("href", "")
            full = urljoin(base_url, href)
            if not full.startswith("http"):
                continue
            if urlparse(full).netloc != base_host:
                continue
            links.append(full)
            if len(links) >= limit:
                break
    except Exception:
        return []
    return links


def fetch_page_text_deep(url: str, max_chars: int = 4000) -> str:
    # primary page
    text = fetch_page_text(url, max_chars=max_chars)
    if not text:
        return ""
    # attempt internal links
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept-Language": "zh-CN,zh;q=0.9",
            "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        }
        resp = requests.get(url, headers=headers, timeout=12)
        resp.raise_for_status()
        html = decode_html_bytes(resp.content)
        links = extract_internal_links(url, html, limit=3)
    except Exception:
        links = []
    for link in links:
        t = fetch_page_text(link, max_chars=1500)
        if t:
            text += "\n" + t
    text = clean_text(text)
    if is_gibberish(text):
        return ""
    return text[:max_chars]


def web_search_iter(
    queries: List[str],
    max_results: int = 6,
    fetch_text: bool = True,
    region: str = "cn_only",
):
    results: List[Dict[str, Any]] = []
    if not queries:
        return
    bad = [
        "bing.com/search",
        "m.bing.com/search",
        "cc.bingj.com/cache",
        "bing.com/aclick",
        "bing.com/ck",
        "duckduckgo.com/?q=",
    ]
    region_code = "cn-zh" if region in ("cn_only", "cn_priority") else "wt-wt"

    def run_ddg(ddgs: DDGS, q: str, backend: str, limit: int):
        return list(
            ddgs.text(
                q,
                region=region_code,
                safesearch="off",
                backend=backend,
                max_results=limit,
            )
        )

    backends_used = set()
    with DDGS() as ddgs:
        for q in queries:
            if len(results) >= max_results:
                break
            try:
                remaining = max(1, max_results - len(results))
                # Avoid bing backend; prefer html, fallback lite
                items = run_ddg(ddgs, q, "html", remaining)
                if items:
                    backends_used.add("html")
                if not items:
                    items = run_ddg(ddgs, q, "lite", remaining)
                    if items:
                        backends_used.add("lite")
                logger.info("DDG (%s) results for '%s': %s", "html/lite", q, len(items))
                items_meta = []
                for r in items:
                    if len(items_meta) + len(results) >= max_results:
                        break
                    url = r.get("href") or ""
                    if not url or any(b in url for b in bad):
                        continue
                    if not filter_by_region(url, region):
                        continue
                    items_meta.append(
                        {
                            "query": q,
                            "title": clean_text(r.get("title") or ""),
                            "url": url,
                            "snippet": clean_text(r.get("body") or ""),
                            "content": "",
                        }
                    )
                if not fetch_text:
                    for item in items_meta:
                        if not item["snippet"]:
                            continue
                        results.append(item)
                        yield item
                else:
                    with ThreadPoolExecutor(max_workers=6) as ex:
                        future_map = {ex.submit(fetch_page_text, it["url"]): it for it in items_meta}
                        for fut in as_completed(future_map):
                            item = future_map[fut]
                            try:
                                item["content"] = fut.result() or ""
                            except Exception:
                                item["content"] = ""
                            if not item["content"] and not item["snippet"]:
                                continue
                            results.append(item)
                            yield item
            except Exception as e:
                logger.warning("Web search failed for query '%s': %s", q, e)
    content_hits = sum(1 for r in results if r.get("content"))
    with SEARCH_DIAG_LOCK:
        SEARCH_DIAG.update(
            {
                "engine": "ddgs",
                "backends": ",".join(sorted(list(backends_used))) or "none",
                "region": region_code,
                "queries": queries,
                "total_results": len(results),
                "content_hits": content_hits,
                "ts": int(time.time()),
            }
        )


def web_search(
    queries: List[str],
    max_results: int = 6,
    fetch_text: bool = True,
    region: str = "cn_only",
) -> List[Dict[str, Any]]:
    results: List[Dict[str, Any]] = []
    for item in web_search_iter(queries, max_results=max_results, fetch_text=fetch_text, region=region):
        results.append(item)
    return results




def evaluate_search_sufficiency(
    api_key: str, base_url: str, model: str, user_query: str, results: List[Dict[str, Any]]
) -> bool:
    # Heuristic first
    total_chars = sum(len(r.get("content", "") or "") for r in results)
    if total_chars >= 2000:
        return True
    # LLM judge if enough
    sample = [
        {
            "title": r.get("title", ""),
            "url": r.get("url", ""),
            "snippet": r.get("snippet", ""),
            "content": (r.get("content", "") or "")[:800],
        }
        for r in results[:6]
    ]
    system_prompt = "你是检索质量评估器，只输出JSON。"
    user_prompt = (
        f"用户问题：{user_query}\n\n"
        f"检索结果样本：{json.dumps(sample, ensure_ascii=False)}\n\n"
        "判断信息是否足以回答问题。输出JSON：\n"
        '{ "sufficient": true/false, "reason": "简短原因" }'
    )
    text = call_reasoner(
        api_key=api_key,
        base_url=base_url,
        model=model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=0.2,
        max_tokens=300,
    )
    data = extract_json_block(text) or safe_json_loads(text) or {}
    if not isinstance(data, dict):
        data = {}
    return bool(data.get("sufficient", False))


def _pick_cjk_font() -> str:
    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
        return "STSong-Light"
    except Exception:
        pass
    # try common Windows fonts
    candidates = [
        ("MicrosoftYaHei", "C:/Windows/Fonts/msyh.ttc"),
        ("SimSun", "C:/Windows/Fonts/simsun.ttc"),
        ("SimHei", "C:/Windows/Fonts/simhei.ttf"),
    ]
    for name, path in candidates:
        try:
            if Path(path).exists():
                pdfmetrics.registerFont(TTFont(name, path))
                return name
        except Exception:
            continue
    return "Helvetica"


def _wrap_text(text: str, font: str, size: int, max_width: float) -> List[str]:
    lines: List[str] = []
    for para in (text or "").splitlines():
        if not para.strip():
            lines.append("")
            continue
        current = ""
        for ch in para:
            w = pdfmetrics.stringWidth(current + ch, font, size)
            if w <= max_width:
                current += ch
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def _extract_name_from_text(text: str) -> str:
    if not text:
        return "未署名"
    patterns = [
        r"姓名[:：\s]*([\\u4e00-\\u9fff]{2,4})",
        r"申诉人[:：\s]*([\\u4e00-\\u9fff]{2,4})",
        r"投诉人[:：\s]*([\\u4e00-\\u9fff]{2,4})",
        r"原告[:：\s]*([\\u4e00-\\u9fff]{2,4})",
        r"委托人[:：\s]*([\\u4e00-\\u9fff]{2,4})",
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            return m.group(1)
    return "未署名"


def _safe_filename(name: str) -> str:
    name = re.sub(r"[\\\\/:*?\"<>|]", "_", name)
    name = name.replace(" ", "_")
    return name[:80]


def generate_pdf(title: str, body: str, images: Optional[List[str]] = None, doc_type: Optional[str] = None) -> str:
    doc_type = doc_type or title or "文书"
    name = _extract_name_from_text(body)
    date = time.strftime("%Y%m%d")
    filename = _safe_filename(f"{doc_type}-{name}-{date}.pdf")
    path = EXPORT_DIR / filename
    c = canvas.Canvas(str(path), pagesize=A4)
    width, height = A4
    margin_x = 50
    margin_top = 60
    margin_bottom = 60
    line_gap = 18
    font_main = _pick_cjk_font()
    logo_img = None
    logo_path = WEB_DIR / "assets" / "logo.jpg"
    if not logo_path.exists():
        logo_path = WEB_DIR / "assets" / "logo.png"
    if logo_path.exists():
        try:
            logo_img = ImageReader(str(logo_path))
        except Exception:
            logo_img = None

    def draw_header_footer(page_no: int):
        c.setFont(font_main, 10)
        c.drawString(margin_x, height - 30, "LAW MASTER v1.0")
        title_right = width - margin_x
        if logo_img is not None:
            logo_w = 30
            logo_h = 30
            logo_x = width - margin_x - logo_w
            logo_y = height - 34
            try:
                c.drawImage(logo_img, logo_x, logo_y, width=logo_w, height=logo_h, mask="auto")
                title_right = logo_x - 6
            except Exception:
                pass
        c.drawRightString(title_right, height - 30, title[:40])
        c.line(margin_x, height - 36, width - margin_x, height - 36)
        c.line(margin_x, margin_bottom - 10, width - margin_x, margin_bottom - 10)
        c.setFont(font_main, 9)
        c.drawRightString(width - margin_x, margin_bottom - 24, f"第 {page_no} 页")
        c.drawString(margin_x, margin_bottom - 24, time.strftime("%Y-%m-%d"))

    page_no = 1
    draw_header_footer(page_no)
    y = height - margin_top
    c.setFont(font_main, 16)
    c.drawString(margin_x, y, title)
    y -= 26
    c.setFont(font_main, 11)

    def new_page():
        nonlocal y, page_no
        c.showPage()
        page_no += 1
        draw_header_footer(page_no)
        y = height - margin_top
        c.setFont(font_main, 11)

    # handle image placeholders [[IMAGE:asset_id]]
    parts = re.split(r"\[\[IMAGE:([a-f0-9]+)\]\]", body or "")
    segment_texts: List[str] = []
    segment_images: List[str] = []
    if len(parts) > 1:
        for i, p in enumerate(parts):
            if i % 2 == 0:
                segment_texts.append(p)
            else:
                segment_images.append(p)
    else:
        segment_texts = [body or ""]

    def draw_text_block(text_block: str):
        nonlocal y
        lines = _wrap_text(text_block, font_main, 11, width - margin_x * 2)
        for line in lines:
            if y < margin_bottom + 20:
                new_page()
            if not line:
                y -= line_gap
                continue
            c.drawString(margin_x, y, line)
            y -= line_gap

    def draw_image(asset_id: str):
        nonlocal y
        path = find_asset_path(asset_id)
        if not path or not path.exists():
            return
        try:
            img = ImageReader(str(path))
            iw, ih = img.getSize()
            max_w = width - margin_x * 2
            max_h = 260
            scale = min(max_w / iw, max_h / ih, 1.0)
            w = iw * scale
            h = ih * scale
            if y - h < margin_bottom + 20:
                new_page()
            c.drawImage(img, margin_x, y - h, width=w, height=h)
            y -= h + 12
        except Exception:
            return

    if segment_images:
        for idx, text_block in enumerate(segment_texts):
            draw_text_block(text_block)
            if idx < len(segment_images):
                draw_image(segment_images[idx])
    else:
        draw_text_block(body or "")
        for asset_id in images or []:
            draw_image(asset_id)

    c.save()
    return f"/exports/{filename}"


def rewrite_query(api_key: str, base_url: str, model: str, query: str) -> Dict[str, Any]:
    system_prompt = (
        "你是法律检索改写助手。将用户口语问题改写为检索关键词与向量检索查询。"
        "只输出JSON，不要额外文本。"
    )
    user_prompt = (
        "用户问题：" + query + "\n\n"
        "请输出JSON，格式如下：\n"
        "{\n"
        "  \"keywords_for_search\": [\"关键词1\", \"关键词2\"],\n"
        "  \"query_for_vector_search\": \"用于语义检索的完整查询\"\n"
        "}\n"
        "要求：关键词是法律术语、核心事实与争议点；不要解释。"
    )
    text = call_deepseek(
        api_key=api_key,
        base_url=base_url,
        model=model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=CONFIG.rewrite_temperature,
        max_tokens=256,
    )
    data = extract_json_block(text) or safe_json_loads(text)
    if not data:
        data = {
            "keywords_for_search": fallback_keywords(query),
            "query_for_vector_search": query,
        }
    if "keywords_for_search" not in data:
        data["keywords_for_search"] = fallback_keywords(query)
    if "query_for_vector_search" not in data:
        data["query_for_vector_search"] = query
    return data


def generate_search_keywords(
    api_key: str,
    base_url: str,
    model: str,
    user_query: str,
    memory_context: str = "",
) -> Dict[str, Any]:
    system_prompt = "你是法律检索与网页搜索关键词生成器，只输出JSON。"
    user_prompt = (
        f"用户问题：{user_query}\n"
        f"历史要点：{memory_context or '(无)'}\n\n"
        "请输出JSON：\n"
        "{\n"
        '  "law_keywords": ["法条关键词1","法条关键词2"],\n'
        '  "case_keywords": ["案例关键词1","案例关键词2"],\n'
        '  "web_queries": ["网页查询1","网页查询2","网页查询3"],\n'
        '  "vector_query": "用于语义检索的完整查询"\n'
        "}\n"
        "要求：\n"
        "1) 不要引入无关实体或政策名称。\n"
        "2) law_keywords 偏向法条/权利/违法类型。\n"
        "3) case_keywords 偏向案由/纠纷类型/事实关键词。\n"
        "4) web_queries 用于搜索引擎，尽量具体。\n"
    )
    text = call_deepseek(
        api_key=api_key,
        base_url=base_url,
        model=model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=0.2,
        max_tokens=400,
    )
    data = extract_json_block(text) or safe_json_loads(text) or {}
    if not isinstance(data, dict):
        data = {}
    law_keywords = data.get("law_keywords") or []
    case_keywords = data.get("case_keywords") or []
    web_queries = data.get("web_queries") or []
    vector_query = data.get("vector_query") or data.get("query_for_vector_search") or user_query
    if not law_keywords:
        law_keywords = fallback_keywords(user_query, k=6)
    if not case_keywords:
        case_keywords = law_keywords
    if not web_queries:
        web_queries = heuristic_queries(user_query)
    return {
        "law_keywords": [k.strip() for k in law_keywords if str(k).strip()],
        "case_keywords": [k.strip() for k in case_keywords if str(k).strip()],
        "web_queries": [q.strip() for q in web_queries if str(q).strip()],
        "vector_query": vector_query,
    }


def should_generate_document(text: str) -> Optional[str]:
    if not text:
        return None
    mapping = {
        "申诉": "申诉书",
        "申诉信": "申诉书",
        "投诉": "投诉信",
        "投诉信": "投诉信",
        "律师函": "律师函",
        "起诉": "起诉状",
        "起诉状": "起诉状",
        "文书": "通用文书",
        "说明书": "通用文书",
        "申请书": "通用文书",
    }
    for k, v in mapping.items():
        if k in text:
            return v
    return None


def generate_document_with_model(
    api_key: str,
    base_url: str,
    model: str,
    user_query: str,
    memory_context: str,
    file_text: str,
    file_assets: List[str],
    doc_type: str,
    kb_context: str = "",
    web_context: str = "",
) -> Dict[str, Any]:
    if doc_type not in DOCUMENT_TEMPLATES:
        doc_type = "通用文书"
    tpl = "\n".join([f"- {x}" for x in DOCUMENT_TEMPLATES.get(doc_type, [])])
    system_prompt = (
        "你是法律文书写作助手。只输出JSON。"
        "目标：根据用户需求生成结构清晰、正式、可提交的中文法律文书。"
        "注意：可以在正文中插入 [[IMAGE:asset_id]] 来引用图片证据。"
        "优先引用相关法条与相似案例以增强说服力。"
        "必须将用户事实与法条要件逐条对应，清楚说明对方具体行为如何构成（例如欺诈/违约/侵权）。"
    )
    user_prompt = (
        f"用户需求：{user_query}\n\n"
        f"历史要点：{memory_context or '(无)'}\n\n"
        f"用户补充材料：{file_text}\n\n"
        f"知识库要点（法条/案例）：{kb_context or '(无)'}\n\n"
        f"网页检索摘要：{web_context or '(无)'}\n\n"
        f"可用图片资产ID：{file_assets}\n\n"
        f"文书类型：{doc_type}\n"
        f"文书模板参考：\n{tpl}\n\n"
        "写作要求：\n"
        "1) 在文书正文中先简明叙述用户事件，时间/地点/主体/经过清晰。\n"
        "2) 法律依据必须逐条解释：引用条文后，用通俗语言解释条文含义，并说明与本案事实如何对应。\n"
        "3) 若涉及欺诈/违约/侵权等，请明确构成要件，并对应到对方具体行为。\n"
        "4) 引用案例时，点出案例事实与本案相似点与差异点，增强说服力。\n"
        "5) 请求事项具体可执行（例如退款金额、赔偿项目、期限）。\n\n"
        "请输出JSON：\n"
        "{\n"
        '  "answer": "简要说明已生成文书与下一步建议",\n'
        '  "documents": [\n'
        '    {"title": "文书标题", "body": "正文内容(可含[[IMAGE:asset_id]])", "images": ["asset_id"]}\n'
        "  ]\n"
        "}\n"
    )
    text = call_deepseek(
        api_key=api_key,
        base_url=base_url,
        model=model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=0.4,
        max_tokens=CONFIG.answer_max_tokens,
    )
    normalized = normalize_model_json(text)
    data = extract_json_block(normalized) or safe_json_loads(normalized) or {}
    if not isinstance(data, dict):
        data = {}
    answer = data.get("answer") or extract_answer_field(normalized) or "已生成文书草稿。"
    documents = data.get("documents") or []
    return {"answer": sanitize_answer_text(answer), "documents": documents}


def analyze_case_with_model(
    api_key: str,
    base_url: str,
    model: str,
    user_query: str,
    case_title: str,
    case_content: str,
) -> str:
    system_prompt = "你是法律案例分析助手，输出简洁要点。"
    user_prompt = (
        f"用户问题：{user_query}\n\n"
        f"案例标题：{case_title}\n"
        f"案例内容：{truncate(case_content, 2000)}\n\n"
        "请输出：\n"
        "1) 与用户事件的相似点\n"
        "2) 可借鉴的处理思路/证据要点\n"
        "3) 注意差异与风险提示\n"
        "要求：简洁、分点，不要编造事实。"
    )
    text = call_deepseek(
        api_key=api_key,
        base_url=base_url,
        model=model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=0.4,
        max_tokens=500,
    )
    return sanitize_answer_text(text)


def batch_analyze_cases(
    api_key: str,
    base_url: str,
    model: str,
    user_query: str,
    cases: List[Dict[str, Any]],
    max_cases: int = 3,
) -> List[Dict[str, Any]]:
    if not cases:
        return []
    selected_raw = cases[: max(1, max_cases)]
    selected: List[Dict[str, Any]] = []
    for idx, item in enumerate(selected_raw, start=1):
        row = dict(item)
        row["case_no"] = row.get("case_no") or idx
        selected.append(row)

    def run_case(item: Dict[str, Any]):
        title = item.get("title") or item.get("article_number") or "案例"
        content = item.get("content") or ""
        analysis = analyze_case_with_model(
            api_key=api_key,
            base_url=base_url,
            model=model,
            user_query=user_query,
            case_title=title,
            case_content=content,
        )
        return {
            "case_no": item.get("case_no"),
            "title": title,
            "analysis": analysis,
        }

    workers = min(4, max(1, len(selected)))
    with ThreadPoolExecutor(max_workers=workers) as ex:
        return list(ex.map(run_case, selected))


def generate_mindmap(
    api_key: str,
    base_url: str,
    model: str,
    user_query: str,
    final_answer: str,
    plan_detail: Dict[str, Any],
    kb_brief: Optional[List[Dict[str, Any]]] = None,
    case_brief: Optional[List[Dict[str, Any]]] = None,
    web_brief: Optional[List[Dict[str, Any]]] = None,
    action: str = "",
) -> Dict[str, Any]:
    system_prompt = "你是思维导图生成器。只输出JSON。"
    user_prompt = (
        f"用户问题：{user_query}\n"
        f"最终回答摘要：{truncate(final_answer, 800)}\n"
        f"探索计划：{json.dumps(plan_detail, ensure_ascii=False)}\n"
        f"法条摘要：{json.dumps(kb_brief or [], ensure_ascii=False)}\n"
        f"案例摘要：{json.dumps(case_brief or [], ensure_ascii=False)}\n"
        f"网页摘要：{json.dumps(web_brief or [], ensure_ascii=False)}\n"
        f"动作：{action}\n\n"
        "请输出适合可视化的思维导图JSON，结构如下：\n"
        "{\n"
        '  "nodes": [\n'
        '    {"id":"root","title":"目标","summary":"...","detail":"...","group":"goal"},\n'
        '    {"id":"n1","title":"步骤/结论","summary":"...","detail":"...","group":"plan"},\n'
        '    {"id":"kb1","title":"法条要点","summary":"...","detail":"...","group":"kb"},\n'
        '    {"id":"case1","title":"案例要点","summary":"...","detail":"...","group":"case"},\n'
        '    {"id":"web1","title":"网页要点","summary":"...","detail":"...","group":"web"}\n'
        "  ],\n"
        '  "edges": [ {"from":"root","to":"n1"} ]\n'
        "}\n"
        "要求：\n"
        "1) 节点数量 6-12 个，可有多层结构，不要固定三步。\n"
        "2) 必须包含“事件陈述/事实梳理”节点，清楚概括用户事件、关键事实与争议点，并给出简要法律分析。\n"
        "3) summary 精炼，detail 可稍长，可包含关键事实与分析。\n"
        "4) group 只能为 goal/plan/kb/case/web/action。\n"
    )
    text = call_deepseek(
        api_key=api_key,
        base_url=base_url,
        model=model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=0.3,
        max_tokens=800,
    )
    data = extract_json_block(text) or safe_json_loads(text) or {}
    if not isinstance(data, dict):
        return {}
    nodes = data.get("nodes") or []
    edges = data.get("edges") or []
    if not nodes:
        return {}
    return {"nodes": nodes, "edges": edges}


def build_fallback_mindmap(
    plan_detail: Dict[str, Any],
    kb_brief: Optional[List[Dict[str, Any]]] = None,
    case_brief: Optional[List[Dict[str, Any]]] = None,
    web_brief: Optional[List[Dict[str, Any]]] = None,
    action: str = "",
    user_query: str = "",
) -> Dict[str, Any]:
    nodes = []
    edges = []
    goal = (plan_detail or {}).get("goal") or "目标"
    nodes.append({"id": "root", "title": "目标", "summary": goal, "detail": goal, "group": "goal"})

    if user_query:
        nodes.append(
            {
                "id": "fact",
                "title": "事件陈述",
                "summary": truncate(user_query, 80),
                "detail": truncate(user_query, 240),
                "group": "plan",
            }
        )
        edges.append({"from": "root", "to": "fact"})

    steps = (plan_detail or {}).get("steps") or []
    for i, step in enumerate(steps):
        sid = f"s{i+1}"
        title = step.get("step") or f"步骤{i+1}"
        detail = step.get("detail") or ""
        nodes.append(
            {
                "id": sid,
                "title": title[:24],
                "summary": truncate(detail or title, 80),
                "detail": truncate(detail or title, 240),
                "group": "plan",
            }
        )
        edges.append({"from": "root", "to": sid})
        subs = step.get("subtasks") or []
        for j, sub in enumerate(subs[:4]):
            cid = f"{sid}_c{j+1}"
            nodes.append(
                {
                    "id": cid,
                    "title": truncate(sub, 20),
                    "summary": truncate(sub, 60),
                    "detail": truncate(sub, 160),
                    "group": "plan",
                }
            )
            edges.append({"from": sid, "to": cid})

    if kb_brief:
        kb_titles = [k.get("title") or "法条" for k in kb_brief][:4]
        kb_summary = "、".join(kb_titles)
        nodes.append(
            {
                "id": "kb",
                "title": "法条要点",
                "summary": truncate(kb_summary, 80),
                "detail": truncate(json.dumps(kb_brief, ensure_ascii=False), 400),
                "group": "kb",
            }
        )
        edges.append({"from": "root", "to": "kb"})

    if case_brief:
        case_titles = [k.get("title") or "案例" for k in case_brief][:4]
        case_summary = "、".join(case_titles)
        nodes.append(
            {
                "id": "case",
                "title": "案例要点",
                "summary": truncate(case_summary, 80),
                "detail": truncate(json.dumps(case_brief, ensure_ascii=False), 400),
                "group": "case",
            }
        )
        edges.append({"from": "root", "to": "case"})

    if web_brief:
        web_titles = [k.get("title") or "网页" for k in web_brief][:4]
        web_summary = "、".join(web_titles)
        nodes.append(
            {
                "id": "web",
                "title": "网页要点",
                "summary": truncate(web_summary, 80),
                "detail": truncate(json.dumps(web_brief, ensure_ascii=False), 400),
                "group": "web",
            }
        )
        edges.append({"from": "root", "to": "web"})

    if action:
        nodes.append(
            {
                "id": "action",
                "title": "动作",
                "summary": action,
                "detail": action,
                "group": "action",
            }
        )
        edges.append({"from": "root", "to": "action"})

    return {"nodes": nodes, "edges": edges}


def generate_clarifying_questions(
    api_key: str,
    base_url: str,
    model: str,
    user_query: str,
    memory_context: str,
) -> List[str]:
    system_prompt = "你是法律咨询澄清助手。只输出JSON。"
    user_prompt = (
        f"用户问题：{user_query}\n"
        f"历史要点：{memory_context or '(无)'}\n\n"
        "请给出2-4个澄清问题，像律师提问，帮助明确事实与目标。\n"
        '输出JSON：{ "questions": ["问题1","问题2","问题3"] }'
    )
    text = call_reasoner(
        api_key=api_key,
        base_url=base_url,
        model=model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=0.2,
        max_tokens=240,
    )
    data = extract_json_block(text) or safe_json_loads(text) or {}
    if not isinstance(data, dict):
        return []
    qs = data.get("questions") or []
    cleaned = [str(q).strip() for q in qs if str(q).strip()]
    return cleaned


def generate_quick_clarify(
    api_key: str,
    base_url: str,
    model: str,
    user_query: str,
    memory_context: str,
) -> List[str]:
    system_prompt = "你是法律咨询澄清助手。只输出JSON。"
    user_prompt = (
        f"用户问题：{user_query}\n"
        f"历史要点：{memory_context or '(无)'}\n\n"
        "快速给出2-4个澄清问题，像律师提问，聚焦事实、证据、诉求与时间线。\n"
        '输出JSON：{ "questions": ["问题1","问题2","问题3"] }'
    )
    text = call_deepseek(
        api_key=api_key,
        base_url=base_url,
        model=model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=0.2,
        max_tokens=200,
    )
    data = extract_json_block(text) or safe_json_loads(text) or {}
    if not isinstance(data, dict):
        return []
    qs = data.get("questions") or []
    return [str(q).strip() for q in qs if str(q).strip()]


def generate_answer(
    api_key: str,
    base_url: str,
    model: str,
    query: str,
    context: str,
    case_context: str,
    history: List[Dict[str, str]],
    file_text: str = "",
    memory_context: str = "",
) -> str:
    system_prompt = (
        "你是LAW MASTER v1.0，一个严谨、清晰的中文法律咨询助手。"
        "必须优先依据提供的法条回答，不要编造不存在的法条。"
        "如有相似案例，请引用并说明案例要点以增强说服力。"
        "需要将用户事实与法条要件逐条对应，明确说明对方具体行为如何构成相关法律要件。"
        "最后请提示：仅供参考，不构成正式法律意见。"
    )

    pairs = []
    for msg in history[-8:]:
        if msg.get("role") == "user":
            pairs.append([msg.get("content", ""), ""])
        elif msg.get("role") == "assistant" and pairs:
            pairs[-1][1] = msg.get("content", "")
    history_text = "\n".join([f"用户：{u}\n助手：{a}" for u, a in pairs[-4:]])

    user_prompt = (
        "【对话历史】\n" + (history_text or "(无)") + "\n\n"
        "【历史要点】\n" + (memory_context or "(无)") + "\n\n"
        "【法条上下文】\n" + (context or "(未检索到相关法条)") + "\n\n"
        "【案例上下文】\n" + (case_context or "(未检索到相关案例)") + "\n\n"
        "【用户补充材料】\n" + (file_text or "(无)") + "\n\n"
        "【用户问题】\n" + query + "\n\n"
        "请用清晰中文输出：\n"
        "1. 结论（1-3句）\n"
        "2. 法条依据（引用具体条文编号，并用通俗话解释条文含义）\n"
        "3. 分析（围绕事实与争议点，说明事实如何满足条文要件）\n"
        "4. 建议（可操作步骤）\n"
        "5. 风险提示（如证据不足等）\n"
    )

    return call_deepseek(
        api_key=api_key,
        base_url=base_url,
        model=model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=CONFIG.answer_temperature,
        max_tokens=CONFIG.answer_max_tokens,
    )


# -----------------------------
# Retrieval
# -----------------------------


def retrieve_docs(
    docs: List[Dict[str, Any]],
    vectorizer: TfidfVectorizer,
    matrix,
    query_vector: str,
    keywords: List[str],
    top_k: int,
    keyword_weight: float,
    vector_weight: float,
):
    if not docs or vectorizer is None or matrix is None:
        return []
    vector_scores = cosine_similarity(vectorizer.transform([query_vector]), matrix).flatten()

    keyword_scores = np.zeros(len(docs), dtype=np.float32)
    keyword_set = {k.strip() for k in (keywords or []) if k.strip()}
    if keyword_set:
        for i, doc in enumerate(docs):
            doc_keywords = set(doc.get("keywords", []) or [])
            if doc_keywords:
                hits = keyword_set.intersection(doc_keywords)
                if hits:
                    keyword_scores[i] = float(len(hits))
            else:
                # fallback: keyword match in title/content (useful for case docs)
                blob = f"{doc.get('title','')} {doc.get('content','')}"
                hits = sum(1 for k in keyword_set if k and k in blob)
                if hits:
                    keyword_scores[i] = float(hits)
            # extra boost for case title matches
            if doc.get("doc_type") == "case" and keyword_set:
                title = doc.get("title", "")
                if title:
                    title_hits = sum(1 for k in keyword_set if k and k in title)
                    if title_hits:
                        keyword_scores[i] += float(title_hits) * 0.5

    if keyword_scores.max() > 0:
        keyword_scores = keyword_scores / keyword_scores.max()
    if vector_scores.max() > 0:
        vector_scores = vector_scores / vector_scores.max()

    fused = keyword_weight * keyword_scores + vector_weight * vector_scores
    top_idx = np.argsort(fused)[::-1][:top_k]

    results = []
    for idx in top_idx:
        doc = docs[int(idx)]
        results.append(
            {
                "id": doc.get("id", idx),
                "doc_type": doc.get("doc_type") or "law",
                "title": doc.get("title", ""),
                "article_number": doc.get("article_number", ""),
                "content": doc.get("content", ""),
                "source": doc.get("source", ""),
                "page_start": doc.get("page_start"),
                "page_end": doc.get("page_end"),
                "summary": doc.get("summary", ""),
                "score": float(fused[idx]),
                "score_vector": float(vector_scores[idx]),
                "score_keyword": float(keyword_scores[idx]),
            }
        )
    return results


def build_context(docs: List[Dict[str, Any]]) -> str:
    parts = []
    for i, d in enumerate(docs):
        doc_type = d.get("doc_type") or "law"
        content = d.get("content", "")
        source = d.get("source", "")
        if doc_type == "case":
            title = d.get("title") or d.get("article_number") or f"案例_{d.get('id', i)}"
            parts.append(f"[案例{i+1}] {title}\n来源：{source}\n{content}")
        else:
            article = d.get("article_number") or f"条文_{d.get('id', i)}"
            parts.append(f"[{i+1}] {article}\n来源：{source}\n{content}")
    return "\n\n".join(parts)


def update_extension_kb(entries: List[Dict[str, Any]]) -> int:
    global DOCS, EXT_DOCS, _next_id
    if not entries:
        return 0
    added = 0
    with EXT_KB_PATH.open("a", encoding="utf-8") as f:
        for i, e in enumerate(entries):
            content = (e.get("content") or "").strip()
            if not content:
                continue
            doc = {
                "id": _next_id,
                "article_number": e.get("article_number")
                or f"扩展知识库-{int(time.time())}-{_next_id}",
                "source": e.get("source") or e.get("url") or "web",
                "content": content,
                "keywords": e.get("keywords") or [],
                "hypothetical_questions": [],
                "doc_type": "law",
            }
            f.write(json.dumps(doc, ensure_ascii=False) + "\n")
            EXT_DOCS.append(doc)
            DOCS.append(doc)
            _next_id += 1
            added += 1
    if added:
        rebuild_index()
    return added


# -----------------------------
# File parsing
# -----------------------------


def extract_text_from_file(path: Path, ocr_engine: str = "paddle") -> Tuple[str, str]:
    suffix = path.suffix.lower()
    if suffix in [".txt", ".md", ".json", ".jsonl"]:
        text = path.read_text(encoding="utf-8", errors="ignore")
        return clean_ocr_text(text), f"已读取 {path.name} (文本)"
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except Exception:
            return "", "缺少依赖 pypdf"
        reader = PdfReader(str(path))
        pdf_text = "\n".join([page.extract_text() or "" for page in reader.pages])
        return clean_ocr_text(pdf_text), f"已读取 {path.name} (PDF)"
    if suffix == ".docx":
        try:
            import docx
        except Exception:
            return "", "缺少依赖 python-docx"
        doc = docx.Document(str(path))
        doc_text = "\n".join([para.text for para in doc.paragraphs])
        return clean_ocr_text(doc_text), f"已读取 {path.name} (DOCX)"
    if suffix in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"]:
        text, err, engine_used = ocr_image_file(path, ocr_engine)
        if err:
            return "", err
        engine_note = f" ({engine_used})" if engine_used else ""
        return clean_ocr_text(text or ""), f"已OCR识别 {path.name}{engine_note}"
    return "", f"不支持的格式：{path.name}"


# -----------------------------
# Voice (optional)
# -----------------------------


_whisper_model = None
_ocr_model = None
DETECTED_TESSERACT_CMD = ""


def get_tesseract_cmd() -> str:
    env_cmd = os.getenv("TESSERACT_CMD", "").strip()
    if env_cmd:
        if Path(env_cmd).exists():
            return env_cmd
        which = shutil.which(env_cmd)
        if which:
            return which

    candidates = [
        PROJECT_ROOT / "tools" / "tesseract" / "tesseract.exe",
        PROJECT_ROOT / "tools" / "Tesseract-OCR" / "tesseract.exe",
        r"C:\\Program Files\\Tesseract-OCR\\tesseract.exe",
        r"C:\\Program Files (x86)\\Tesseract-OCR\\tesseract.exe",
    ]
    for cand in candidates:
        if Path(cand).exists():
            return str(cand)

    which = shutil.which("tesseract")
    return which or ""


def ensure_tesseract_env() -> str:
    global DETECTED_TESSERACT_CMD
    if not os.getenv("TESSERACT_CMD"):
        cmd = get_tesseract_cmd()
        if cmd:
            os.environ["TESSERACT_CMD"] = cmd
    DETECTED_TESSERACT_CMD = os.getenv("TESSERACT_CMD", "") or get_tesseract_cmd()
    return DETECTED_TESSERACT_CMD


def import_paddleocr_with_torch_stub():
    import types
    import sys as _sys
    from importlib.machinery import ModuleSpec

    for key in [
        "torch",
        "torch.cuda",
        "torch.version",
        "torch.nn",
        "torch.multiprocessing",
        "torch.distributed",
    ]:
        if key in _sys.modules:
            del _sys.modules[key]

    torch_stub = types.ModuleType("torch")
    torch_stub.__version__ = "0.0"
    torch_stub.__spec__ = ModuleSpec(name="torch", loader=None)
    torch_stub.__path__ = []

    cuda_mod = types.ModuleType("torch.cuda")
    cuda_mod.__spec__ = ModuleSpec(name="torch.cuda", loader=None)
    cuda_mod.is_available = lambda: False
    cuda_mod.get_device_capability = lambda: (0, 0)

    version_mod = types.ModuleType("torch.version")
    version_mod.__spec__ = ModuleSpec(name="torch.version", loader=None)
    version_mod.cuda = None

    nn_mod = types.ModuleType("torch.nn")
    nn_mod.__spec__ = ModuleSpec(name="torch.nn", loader=None)

    mp_mod = types.ModuleType("torch.multiprocessing")
    mp_mod.__spec__ = ModuleSpec(name="torch.multiprocessing", loader=None)
    try:
        import multiprocessing as _mp

        mp_mod.get_start_method = _mp.get_start_method
        mp_mod.set_start_method = _mp.set_start_method
        mp_mod.cpu_count = _mp.cpu_count
    except Exception:
        mp_mod.get_start_method = lambda *args, **kwargs: "spawn"
        mp_mod.set_start_method = lambda *args, **kwargs: None
        mp_mod.cpu_count = lambda: 1

    dist_mod = types.ModuleType("torch.distributed")
    dist_mod.__spec__ = ModuleSpec(name="torch.distributed", loader=None)
    dist_mod.is_available = lambda: False
    dist_mod.is_initialized = lambda: False
    dist_mod.init_process_group = lambda *args, **kwargs: None
    dist_mod.destroy_process_group = lambda *args, **kwargs: None

    torch_stub.cuda = cuda_mod
    torch_stub.version = version_mod
    torch_stub.nn = nn_mod
    torch_stub.multiprocessing = mp_mod
    torch_stub.distributed = dist_mod

    _sys.modules["torch"] = torch_stub
    _sys.modules["torch.cuda"] = cuda_mod
    _sys.modules["torch.version"] = version_mod
    _sys.modules["torch.nn"] = nn_mod
    _sys.modules["torch.multiprocessing"] = mp_mod
    _sys.modules["torch.distributed"] = dist_mod

    import paddle
    from paddleocr import PaddleOCR

    return paddle, PaddleOCR


def transcribe_audio_file(path: Path, model_size: str = "base") -> Tuple[Optional[str], Optional[str]]:
    global _whisper_model
    try:
        from faster_whisper import WhisperModel
    except Exception:
        return None, "缺少依赖 faster-whisper，无法语音识别。"

    if _whisper_model is None:
        _whisper_model = WhisperModel(model_size, device="cpu", compute_type="int8")

    try:
        segments, _info = _whisper_model.transcribe(str(path), language="zh")
        text = "".join([seg.text for seg in segments])
        return text.strip(), None
    except Exception as e:
        return None, f"语音识别失败：{e}"


def ocr_image_file(path: Path, engine: str = "paddle") -> Tuple[Optional[str], Optional[str], str]:
    global _ocr_model
    engine = (engine or "auto").strip().lower()
    if engine not in {"auto", "tesseract", "paddle"}:
        engine = "auto"

    last_err = None

    # 1) Try Tesseract
    if engine in {"auto", "tesseract"}:
        try:
            import pytesseract
            from PIL import Image

            tcmd = get_tesseract_cmd()
            if tcmd:
                pytesseract.pytesseract.tesseract_cmd = tcmd
                lang = os.getenv("TESSERACT_LANG", "chi_sim+eng")
                text = pytesseract.image_to_string(Image.open(path), lang=lang)
                text = clean_text(text)
                if text:
                    return text, None, "tesseract"
                last_err = "tesseract 识别未返回文字"
            else:
                last_err = "未检测到 tesseract 可执行文件，请先安装 Tesseract-OCR。"
        except Exception as e:
            last_err = f"tesseract 识别失败: {e}"
            logger.warning("Tesseract OCR failed: %s", e)

        if engine == "tesseract":
            return None, last_err or "tesseract 识别失败", "tesseract"

    # 2) Fallback to PaddleOCR
    try:
        import paddle
        from paddleocr import PaddleOCR
    except Exception as e:
        msg = str(e)
        if "shm.dll" in msg or "WinError 127" in msg:
            try:
                paddle, PaddleOCR = import_paddleocr_with_torch_stub()
            except Exception as e2:
                return None, f"无法导入 PaddleOCR(疑似Torch依赖问题): {e2}", "paddle"
        else:
            return None, f"无法导入 PaddleOCR: {e}", "paddle"

    if _ocr_model is None:
        try:
            paddle.set_flags(
                {
                    "FLAGS_enable_pir_api": False,
                    "FLAGS_enable_pir_in_executor": False,
                    "FLAGS_use_mkldnn": False,
                    "FLAGS_enable_mkldnn": False,
                    "FLAGS_enable_onednn": False,
                }
            )
        except Exception:
            pass
        try:
            import inspect

            params = inspect.signature(PaddleOCR).parameters
            kwargs = {"lang": "ch"}
            if "use_angle_cls" in params:
                kwargs["use_angle_cls"] = True
            elif "use_textline_orientation" in params:
                kwargs["use_textline_orientation"] = True
            if "use_gpu" in params:
                kwargs["use_gpu"] = False
            if "use_mkldnn" in params:
                kwargs["use_mkldnn"] = False
            _ocr_model = PaddleOCR(**kwargs)
        except Exception as e:
            return None, f"PaddleOCR 初始化失败: {e}", "paddle"

    try:
        try:
            result = _ocr_model.ocr(str(path), cls=True)
        except TypeError as e:
            if "cls" in str(e):
                result = _ocr_model.ocr(str(path))
            else:
                raise
        lines = []
        for page in result:
            for item in page:
                text = item[1][0]
                if text:
                    lines.append(text)
        return "\n".join(lines), None, "paddle"
    except Exception as e:
        msg = str(e)
        if "ConvertPirAttribute2RuntimeAttribute" in msg or "onednn" in msg:
            return None, "OCR失败：PaddlePaddle/OneDNN 兼容问题，请尝试重启服务或降级 paddlepaddle 与 paddleocr 版本。", "paddle"
        return None, f"OCR失败：{e}", "paddle"


# -----------------------------
# Global init
# -----------------------------
load_dotenv()
os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
os.environ.setdefault("PADDLE_PDX_EAGER_INIT", "False")
ensure_tesseract_env()
DEFAULT_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")
DEFAULT_BASE_URL = os.getenv("DEEPSEEK_BASE_URL", CONFIG.base_url)
DEFAULT_MODEL = os.getenv("DEEPSEEK_MODEL", CONFIG.model)

logger.info("Loading corpus...")
BASE_DOCS = load_corpus(CORPUS_PATH)
EXT_DOCS = load_extension_corpus(EXT_KB_PATH)
DOCS = BASE_DOCS + EXT_DOCS
logger.info("Corpus loaded: %s (base=%s, ext=%s)", len(DOCS), len(BASE_DOCS), len(EXT_DOCS))

logger.info("Loading case corpus...")
CASE_DOCS = load_case_corpus(CASE_KB_PATH)
logger.info("Case corpus loaded: %s (path=%s)", len(CASE_DOCS), CASE_KB_PATH)

logger.info("Building / loading index...")
VECTORIZER, MATRIX, _DOC_TEXTS = load_or_build_index(
    DOCS, CORPUS_PATH, extra_key="_with_ext" + get_ext_cache_key()
)
CASE_VECTORIZER, CASE_MATRIX, _CASE_DOC_TEXTS = (None, None, [])
if CASE_DOCS:
    CASE_VECTORIZER, CASE_MATRIX, _CASE_DOC_TEXTS = load_or_build_index(
        CASE_DOCS, CASE_KB_PATH, extra_key="_case"
    )
logger.info("Index ready")

_max_id = max([d.get("id", -1) for d in DOCS], default=-1)
_next_id = _max_id + 1


# -----------------------------
# FastAPI
# -----------------------------

app = FastAPI(title=CONFIG.name)

if WEB_DIR.exists():
    app.mount("/assets", StaticFiles(directory=WEB_DIR / "assets"), name="assets")
if EXPORT_DIR.exists():
    app.mount("/exports", StaticFiles(directory=EXPORT_DIR), name="exports")


@app.get("/api/search_diagnostics")
def search_diagnostics():
    with SEARCH_DIAG_LOCK:
        return SEARCH_DIAG.copy()


@app.post("/api/generate_document")
async def generate_document(payload: Dict[str, Any] = Body(...)):
    api_key = payload.get("api_key") or DEFAULT_API_KEY
    base_url = payload.get("base_url") or DEFAULT_BASE_URL
    model = payload.get("model") or DEFAULT_MODEL
    session_id = payload.get("session_id") or "default"
    user_query = payload.get("query") or ""
    file_text = clean_ocr_text(payload.get("file_text") or "")
    file_assets = payload.get("file_assets") or []
    doc_type = payload.get("doc_type") or "通用文书"
    title = payload.get("title") or ""
    body = payload.get("body") or ""

    if not api_key:
        return JSONResponse({"error": "缺少 DeepSeek API Key"}, status_code=400)
    if not body and not user_query:
        return JSONResponse({"error": "缺少文书需求"}, status_code=400)

    memory_context = build_memory_context(session_id)

    if body:
        url = generate_pdf(title or doc_type, body, images=file_assets, doc_type=doc_type)
        return {"downloads": [{"title": title or doc_type, "url": url}], "documents": [{"title": title or doc_type, "body": body, "images": file_assets}]}

    # Optional: enrich with KB context
    kb_context = ""
    try:
        rewrite = rewrite_query(api_key, base_url, model, user_query)
        rag_keywords = rewrite.get("keywords_for_search", [])
        query_vector = rewrite.get("query_for_vector_search", user_query)
        law_results = retrieve_docs(
            DOCS,
            VECTORIZER,
            MATRIX,
            query_vector=query_vector,
            keywords=rag_keywords,
            top_k=5,
            keyword_weight=CONFIG.keyword_weight,
            vector_weight=CONFIG.vector_weight,
        )
        case_results = retrieve_docs(
            CASE_DOCS,
            CASE_VECTORIZER,
            CASE_MATRIX,
            query_vector=query_vector,
            keywords=rag_keywords,
            top_k=5,
            keyword_weight=CONFIG.keyword_weight,
            vector_weight=CONFIG.vector_weight,
        )
        kb_context = truncate(build_context(law_results), 2000)
        case_context = truncate(build_context(case_results), 2000)
        kb_context = (kb_context + "\n\n" + case_context).strip()
    except Exception:
        kb_context = ""

    doc_pack = generate_document_with_model(
        api_key=api_key,
        base_url=base_url,
        model=model,
        user_query=user_query,
        memory_context=memory_context,
        file_text=file_text,
        file_assets=file_assets,
        doc_type=doc_type,
        kb_context=kb_context,
    )
    documents = doc_pack.get("documents") or []
    downloads = []
    for doc in documents:
        dtitle = doc.get("title") or doc_type
        dbody = doc.get("body") or ""
        images = doc.get("images") or file_assets
        url = generate_pdf(dtitle, dbody, images=images, doc_type=doc_type)
        downloads.append({"title": dtitle, "url": url})

    doc_titles = [d.get("title") or doc_type for d in documents]
    for entry in extract_sparse_memory(user_query, doc_pack.get("answer", ""), doc_titles=doc_titles):
        append_memory(session_id, entry)

    return {"answer": doc_pack.get("answer", ""), "documents": documents, "downloads": downloads}


@app.post("/api/case_analysis")
async def case_analysis(payload: Dict[str, Any] = Body(...)):
    api_key = payload.get("api_key") or DEFAULT_API_KEY
    base_url = payload.get("base_url") or DEFAULT_BASE_URL
    model = payload.get("model") or DEFAULT_MODEL
    user_query = payload.get("user_query") or ""
    cases = payload.get("cases") or []
    if not api_key:
        return JSONResponse({"error": "缺少 DeepSeek API Key"}, status_code=400)
    if not cases:
        return {"analyses": []}
    results = batch_analyze_cases(
        api_key=api_key,
        base_url=base_url,
        model=model,
        user_query=user_query,
        cases=cases,
        max_cases=len(cases),
    )
    return {"analyses": results}


@app.post("/api/doc_refine")
async def doc_refine(payload: Dict[str, Any] = Body(...)):
    api_key = payload.get("api_key") or DEFAULT_API_KEY
    base_url = payload.get("base_url") or DEFAULT_BASE_URL
    model = payload.get("model") or DEFAULT_MODEL
    session_id = payload.get("session_id") or "default"
    text = payload.get("text") or ""
    instruction = payload.get("instruction") or ""
    doc_type = payload.get("doc_type") or "通用文书"

    if not api_key:
        return JSONResponse({"error": "缺少 DeepSeek API Key"}, status_code=400)
    if not text:
        return JSONResponse({"error": "文书内容为空"}, status_code=400)

    memory_context = build_memory_context(session_id)
    system_prompt = "你是法律文书润色助手。保持结构清晰、正式、逻辑严谨。"
    user_prompt = (
        f"文书类型：{doc_type}\n"
        f"历史要点：{memory_context or '(无)'}\n"
        f"用户指令：{instruction}\n\n"
        f"当前文书：\n{text}\n\n"
        "请输出润色后的完整文书正文，不要额外解释。"
    )
    refined = call_deepseek(
        api_key=api_key,
        base_url=base_url,
        model=model,
        messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
        temperature=0.3,
        max_tokens=CONFIG.answer_max_tokens,
    )
    return {"text": sanitize_answer_text(refined)}


@app.post("/api/session_title")
async def session_title(payload: Dict[str, Any] = Body(...)):
    api_key = payload.get("api_key") or DEFAULT_API_KEY
    base_url = payload.get("base_url") or DEFAULT_BASE_URL
    model = payload.get("model") or DEFAULT_MODEL
    user_query = clean_text(payload.get("user_query") or "")
    assistant_text = clean_text(payload.get("assistant_text") or "")

    if not api_key:
        return JSONResponse({"error": "缺少 DeepSeek API Key"}, status_code=400)
    if not user_query:
        return JSONResponse({"error": "缺少 user_query"}, status_code=400)

    system_prompt = (
        "你是会话命名助手。"
        "请输出一个简短中文标题，4到12字。"
        "只输出标题，不要解释、不要编号、不要标点。"
        "不要以“问题/咨询/结论/摘要”结尾。"
    )
    user_prompt = (
        f"用户问题：{user_query}\n"
        f"助手回复（可选）：{assistant_text[:220] if assistant_text else '无'}\n"
        "请给出会话标题。"
    )

    title = ""
    try:
        raw = call_deepseek(
            api_key=api_key,
            base_url=base_url,
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2,
            max_tokens=48,
        )
        title = sanitize_session_title(raw)
    except Exception as e:
        logger.warning("Session title generation failed: %s", e)

    if not title:
        title = fallback_session_title_from_query(user_query)
    return {"title": title}


@app.get("/", response_class=HTMLResponse)
def index():
    html_path = WEB_DIR / "index.html"
    if not html_path.exists():
        return HTMLResponse("Missing web/index.html", status_code=500)
    return HTMLResponse(html_path.read_text(encoding="utf-8"))


@app.get("/api/ocr_info")
def ocr_info():
    import sys
    from importlib.util import find_spec
    try:
        from importlib.metadata import version, PackageNotFoundError
    except Exception:  # pragma: no cover
        version = None
        PackageNotFoundError = Exception
    cmd = ensure_tesseract_env()
    info = {
        "tesseract_cmd": cmd,
        "has_tesseract": bool(cmd),
        "python": sys.executable,
        "paddle": None,
        "paddleocr": None,
        "paddleocr_installed": False,
        "paddle_error": "",
        "paddleocr_error": "",
    }
    try:
        if version:
            info["paddle"] = version("paddlepaddle")
        else:
            import paddle

            info["paddle"] = getattr(paddle, "__version__", "unknown")
    except Exception as e:
        info["paddle_error"] = str(e)
    try:
        if version:
            info["paddleocr"] = version("paddleocr")
            info["paddleocr_installed"] = True
        else:
            info["paddleocr_installed"] = bool(find_spec("paddleocr"))
    except Exception as e:
        info["paddleocr_error"] = str(e)
    return info


@app.post("/api/parse_file")
async def parse_file(file: UploadFile = File(...), ocr_engine: str = Form("browser")):
    try:
        content = await file.read()
        suffix = Path(file.filename).suffix.lower()
        is_image = is_image_ext(suffix)
        if not is_image and file.content_type:
            is_image = file.content_type.startswith("image/")
        if is_image:
            asset_id = save_upload_asset(file.filename, content)
            note = f"已上传图片 {file.filename} (browser)"
            return {"text": "", "note": note, "asset_id": asset_id, "ocr_engine": "browser"}
        tmp_path = CACHE_DIR / f"upload_{int(time.time())}_{file.filename}"
        tmp_path.write_bytes(content)
        try:
            text, note = extract_text_from_file(tmp_path, ocr_engine)
            text = truncate(text, CONFIG.max_file_chars)
            return {"text": text, "note": note}
        finally:
            try:
                tmp_path.unlink(missing_ok=True)
            except Exception:
                pass
    except Exception as e:
        return JSONResponse({"error": f"文件解析失败: {e}"}, status_code=500)


@app.post("/api/transcribe")
async def transcribe(file: UploadFile = File(...)):
    tmp_path = CACHE_DIR / f"audio_{int(time.time())}_{file.filename}"
    content = await file.read()
    tmp_path.write_bytes(content)
    try:
        text, err = transcribe_audio_file(tmp_path)
        if err:
            return JSONResponse({"error": err}, status_code=400)
        return {"text": text}
    finally:
        try:
            tmp_path.unlink(missing_ok=True)
        except Exception:
            pass


@app.post("/api/chat")
async def chat(payload: Dict[str, Any] = Body(...)):
    api_key = payload.get("api_key") or DEFAULT_API_KEY
    base_url = payload.get("base_url") or DEFAULT_BASE_URL
    model = payload.get("model") or DEFAULT_MODEL
    messages = payload.get("messages") or []
    file_text = clean_ocr_text(payload.get("file_text") or "")
    session_id = payload.get("session_id") or "default"
    top_k = int(payload.get("top_k") or CONFIG.top_k)
    keyword_weight = float(payload.get("keyword_weight") or CONFIG.keyword_weight)
    vector_weight = float(payload.get("vector_weight") or CONFIG.vector_weight)

    if not api_key:
        return JSONResponse({"error": "缺少 DeepSeek API Key"}, status_code=400)
    if not messages:
        return JSONResponse({"error": "消息为空"}, status_code=400)

    user_query = ""
    for msg in reversed(messages):
        if msg.get("role") == "user":
            user_query = msg.get("content", "")
            break

    if not user_query:
        return JSONResponse({"error": "未找到用户问题"}, status_code=400)

    rewrite = rewrite_query(api_key, base_url, model, user_query)
    keywords = rewrite.get("keywords_for_search", [])
    query_vector = rewrite.get("query_for_vector_search", user_query)

    law_results = retrieve_docs(
        DOCS,
        VECTORIZER,
        MATRIX,
        query_vector=query_vector,
        keywords=keywords,
        top_k=top_k,
        keyword_weight=keyword_weight,
        vector_weight=vector_weight,
    )
    case_results = retrieve_docs(
        CASE_DOCS,
        CASE_VECTORIZER,
        CASE_MATRIX,
        query_vector=query_vector,
        keywords=keywords,
        top_k=top_k,
        keyword_weight=keyword_weight,
        vector_weight=vector_weight,
    )
    context = truncate(build_context(law_results), CONFIG.max_context_chars)
    case_context = truncate(build_context(case_results), CONFIG.max_context_chars)
    memory_context = build_memory_context(session_id)

    answer = generate_answer(
        api_key=api_key,
        base_url=base_url,
        model=model,
        query=user_query,
        context=context,
        case_context=case_context,
        history=messages,
        file_text=file_text,
        memory_context=memory_context,
    )

    for entry in extract_sparse_memory(user_query, answer):
        append_memory(session_id, entry)

    return {
        "answer": answer,
        "keywords": keywords,
        "results": law_results,
        "law_results": law_results,
        "case_results": case_results,
    }


@app.post("/api/chat_stream")
async def chat_stream(payload: Dict[str, Any] = Body(...)):
    api_key = payload.get("api_key") or DEFAULT_API_KEY
    base_url = payload.get("base_url") or DEFAULT_BASE_URL
    model = payload.get("model") or DEFAULT_MODEL
    messages = payload.get("messages") or []
    file_text = clean_ocr_text(payload.get("file_text") or "")
    session_id = payload.get("session_id") or "default"
    top_k = int(payload.get("top_k") or CONFIG.top_k)
    keyword_weight = float(payload.get("keyword_weight") or CONFIG.keyword_weight)
    vector_weight = float(payload.get("vector_weight") or CONFIG.vector_weight)

    if not api_key:
        return JSONResponse({"error": "缺少 DeepSeek API Key"}, status_code=400)
    if not messages:
        return JSONResponse({"error": "消息为空"}, status_code=400)

    user_query = ""
    for msg in reversed(messages):
        if msg.get("role") == "user":
            user_query = msg.get("content", "")
            break

    if not user_query:
        return JSONResponse({"error": "未找到用户问题"}, status_code=400)

    rewrite = rewrite_query(api_key, base_url, model, user_query)
    keywords = rewrite.get("keywords_for_search", [])
    query_vector = rewrite.get("query_for_vector_search", user_query)

    law_results = retrieve_docs(
        DOCS,
        VECTORIZER,
        MATRIX,
        query_vector=query_vector,
        keywords=keywords,
        top_k=top_k,
        keyword_weight=keyword_weight,
        vector_weight=vector_weight,
    )
    case_results = retrieve_docs(
        CASE_DOCS,
        CASE_VECTORIZER,
        CASE_MATRIX,
        query_vector=query_vector,
        keywords=keywords,
        top_k=top_k,
        keyword_weight=keyword_weight,
        vector_weight=vector_weight,
    )
    context = truncate(build_context(law_results), CONFIG.max_context_chars)
    case_context = truncate(build_context(case_results), CONFIG.max_context_chars)
    memory_context = build_memory_context(session_id)

    answer = generate_answer(
        api_key=api_key,
        base_url=base_url,
        model=model,
        query=user_query,
        context=context,
        case_context=case_context,
        history=messages,
        file_text=file_text,
        memory_context=memory_context,
    )

    def gen():
        yield json.dumps({"type": "status", "label": "模型总结：生成答复"}, ensure_ascii=False) + "\n"
        # naive streaming by chunking final answer
        step = 24
        for i in range(0, len(answer), step):
            chunk = answer[i : i + step]
            yield json.dumps({"type": "delta", "text": chunk}, ensure_ascii=False) + "\n"
            time.sleep(0.02)
        for entry in extract_sparse_memory(user_query, answer):
            append_memory(session_id, entry)
        meta = {
            "type": "meta",
            "keywords": keywords,
            "results": law_results,
            "law_results": law_results,
            "case_results": case_results,
        }
        yield json.dumps(meta, ensure_ascii=False) + "\n"

    return StreamingResponse(gen(), media_type="text/plain")


@app.post("/api/agent_stream")
async def agent_stream(payload: Dict[str, Any] = Body(...)):
    api_key = payload.get("api_key") or DEFAULT_API_KEY
    base_url = payload.get("base_url") or DEFAULT_BASE_URL
    model = payload.get("model") or DEFAULT_MODEL
    reasoning_model = payload.get("reasoning_model") or "deepseek-reasoner"
    site_region = payload.get("site_region") or "cn_priority"
    messages = payload.get("messages") or []
    file_text = clean_ocr_text(payload.get("file_text") or "")
    file_assets = payload.get("file_assets") or []
    session_id = payload.get("session_id") or "default"
    doc_type = payload.get("doc_type") or ""
    force_action = payload.get("force_action") or ""
    force_web_search = bool(payload.get("force_web_search", False))
    top_k = int(payload.get("top_k") or CONFIG.top_k)
    keyword_weight = float(payload.get("keyword_weight") or CONFIG.keyword_weight)
    vector_weight = float(payload.get("vector_weight") or CONFIG.vector_weight)

    if not api_key:
        return JSONResponse({"error": "缺少 DeepSeek API Key"}, status_code=400)
    if not messages:
        return JSONResponse({"error": "消息为空"}, status_code=400)

    user_query = ""
    for msg in reversed(messages):
        if msg.get("role") == "user":
            user_query = msg.get("content", "")
            break

    if not user_query:
        return JSONResponse({"error": "未找到用户问题"}, status_code=400)

    def gen_status(label: str):
        return json.dumps({"type": "status", "label": label}, ensure_ascii=False) + "\n"

    def gen_trace(label: str, detail: str):
        return json.dumps({"type": "trace", "label": label, "detail": detail}, ensure_ascii=False) + "\n"

    def gen():
        progress_val = 0

        def emit_progress(pct: int):
            nonlocal progress_val
            pct = int(max(progress_val, min(100, pct)))
            progress_val = pct
            return json.dumps({"type": "progress", "percent": pct}, ensure_ascii=False) + "\n"

        # non-legal queries: skip agent workflow and respond directly
        if is_non_legal_query(user_query) and not is_legal_query(user_query):
            yield gen_status("模型总结：生成最终答复")
            yield emit_progress(70)
            try:
                system_prompt = "你是生活方式与消费建议助手，回答应简洁、实用、分点给出建议。"
                user_prompt = f"用户问题：{user_query}\n请给出3-6条具体建议，必要时附上挑选原则。"
                text = call_deepseek(
                    api_key=api_key,
                    base_url=base_url,
                    model=model,
                    messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
                    temperature=0.6,
                    max_tokens=600,
                )
            except Exception as e:
                text = f"抱歉，暂时无法生成回答：{e}"
            text = sanitize_answer_text(text)
            for i in range(0, len(text), 24):
                chunk = text[i : i + 24]
                yield json.dumps({"type": "delta", "text": chunk}, ensure_ascii=False) + "\n"
                time.sleep(0.02)
            try:
                state = load_state(session_id)
                if state.get("pending_clarify"):
                    state["pending_clarify"] = False
                    state["pending_for"] = ""
                    state["pending_questions"] = []
                    save_state(session_id, state)
            except Exception:
                pass
            yield emit_progress(100)
            meta = {
                "type": "meta",
                "keywords": [],
                "results": [],
                "law_results": [],
                "case_results": [],
                "search_results": [],
                "highlight_keywords": [],
                "plan_detail": {},
                "downloads": [],
                "documents": [],
                "mindmap": {},
            }
            yield json.dumps(meta, ensure_ascii=False) + "\n"
            return

        # 1) planning
        yield gen_status("模型思考：规划中")
        try:
            memory_context = build_memory_context(session_id)
            state = load_state(session_id)
            pending = bool(state.get("pending_clarify", False))
            if not pending:
                # quick clarify before entering full planning
                try:
                    questions = generate_quick_clarify(
                        api_key=api_key,
                        base_url=base_url,
                        model=model,
                        user_query=user_query,
                        memory_context=memory_context,
                    )
                except Exception:
                    questions = []
                if not questions:
                    questions = [
                        "请简要说明事件背景与关键时间线（发生了什么、何时发生）。",
                        "你的核心目标是什么？例如和解、赔偿、撤销、恢复账号或公开道歉等。",
                        "对方/机构的身份是什么？你目前掌握了哪些证据材料？",
                        "是否需要生成正式文书（申诉书/投诉信/律师函/起诉状）？如需要请说明对象与期限。",
                    ]
                clarify = "为了更准确地完成探索，我需要确认：\n" + "\n".join(
                    [f"{i+1}. {q}" for i, q in enumerate(questions[:4])]
                )
                yield gen_status("模型思考：需要澄清问题")
                for i in range(0, len(clarify), 20):
                    yield json.dumps({"type": "delta", "text": clean_text(clarify[i : i + 20])}, ensure_ascii=False) + "\n"
                    time.sleep(0.02)
                state["pending_clarify"] = True
                state["pending_for"] = user_query
                state["pending_questions"] = questions[:4]
                save_state(session_id, state)
                return
            else:
                # pair user answer with previous questions into memory
                prev_q = state.get("pending_for", "")
                qlist = state.get("pending_questions", []) or []
                if prev_q:
                    append_memory(session_id, {"role": "user", "content": f"初始问题：{prev_q}", "ts": int(time.time())})
                if qlist:
                    append_memory(
                        session_id,
                        {"role": "assistant", "content": "澄清问题：" + "；".join(qlist), "ts": int(time.time())},
                    )
                append_memory(
                    session_id,
                    {"role": "user", "content": "澄清回答：" + user_query, "ts": int(time.time())},
                )
                state["pending_clarify"] = False
                state["pending_for"] = ""
                state["pending_questions"] = []
                state["skip_clarify"] = True
                save_state(session_id, state)

            memory_context = build_memory_context(session_id)
            plan = plan_self_awareness(api_key, base_url, reasoning_model, messages, memory_context=memory_context)
        except Exception as e:
            yield json.dumps({"type": "delta", "text": clean_text(f"规划失败：{e}")}, ensure_ascii=False) + "\n"
            return

        goal = plan.get("goal_analysis", "") or user_query
        plan_steps = normalize_str_list(plan.get("plan", []) or [])
        step_details = normalize_str_list(plan.get("step_details", []) or [])
        subtasks = normalize_subtasks(plan.get("subtasks", []) or [])
        action = (force_action or plan.get("action") or "").strip()
        plan_doc_type = (doc_type or plan.get("document_type") or "").strip()
        explicit_doc = bool(re.search(r"生成|文书|申诉书|投诉信|律师函|起诉状", user_query))
        if not plan_steps:
            plan_steps = ["信息收集与范围界定", "关键信息核验与比对", "形成结论与行动建议"]
            step_details = [
                "梳理目标与适用范围，确定检索方向与关键词。",
                "对比知识库与网页结果，筛选权威信息。",
                "综合结果输出可执行方案与注意事项。",
            ]
        steps = []
        for i, s in enumerate(plan_steps):
            detail = step_details[i] if i < len(step_details) else ""
            subs = subtasks[i] if i < len(subtasks) else []
            if not subs and detail:
                # derive subtasks from detail text
                parts = [p.strip() for p in re.split(r"[；;。\n]+", detail) if p.strip()]
                subs = parts[:3]
            steps.append({"step": s, "detail": detail, "subtasks": subs})
        plan_detail = {
            "goal": goal,
            "steps": steps,
        }
        yield gen_trace("探索计划", "\n".join(plan_steps) or "无")
        yield emit_progress(10)

        state = load_state(session_id)
        if state.get("skip_clarify", False):
            state["skip_clarify"] = False
            save_state(session_id, state)
        else:
            questions = [q for q in (plan.get("clarifying_questions", []) or []) if str(q).strip()]
            if not questions:
                try:
                    questions = generate_clarifying_questions(
                        api_key=api_key,
                        base_url=base_url,
                        model=reasoning_model,
                        user_query=user_query,
                        memory_context=memory_context,
                    )
                except Exception:
                    questions = []
            if not questions:
                questions = [
                    "请简要说明事件背景与关键时间线（发生了什么、何时发生）。",
                    "你的核心目标是什么？例如和解、赔偿、撤销、恢复账号或公开道歉等。",
                    "对方/机构的身份是什么？你目前掌握了哪些证据材料？",
                    "是否需要生成正式文书（申诉书/投诉信/律师函/起诉状）？如需要请说明对象与期限。",
                ]
            clarify = "为了更准确地完成探索，我需要确认：\n" + "\n".join(
                [f"{i+1}. {q}" for i, q in enumerate(questions[:4])]
            )
            yield gen_status("模型思考：需要澄清问题")
            for i in range(0, len(clarify), 20):
                yield json.dumps({"type": "delta", "text": clean_text(clarify[i : i + 20])}, ensure_ascii=False) + "\n"
                time.sleep(0.02)
            state["pending_clarify"] = True
            state["pending_for"] = user_query
            save_state(session_id, state)
            return

        # 2) If document generation action, skip search
        if action == "generate_document" and (force_action == "generate_document" or explicit_doc):
            yield gen_status("模型生成：文书生成中")
            memory_context = build_memory_context(session_id)
            kw_pack = generate_search_keywords(
                api_key=api_key,
                base_url=base_url,
                model=model,
                user_query=user_query,
                memory_context=memory_context,
            )
            law_keywords = kw_pack.get("law_keywords", [])
            case_keywords = kw_pack.get("case_keywords", [])
            rag_keywords = list(dict.fromkeys(law_keywords + case_keywords))
            query_vector = kw_pack.get("vector_query", user_query)
            law_results = retrieve_docs(
                DOCS,
                VECTORIZER,
                MATRIX,
                query_vector=query_vector,
                keywords=law_keywords,
                top_k=top_k,
                keyword_weight=keyword_weight,
                vector_weight=vector_weight,
            )
            case_results = retrieve_docs(
                CASE_DOCS,
                CASE_VECTORIZER,
                CASE_MATRIX,
                query_vector=query_vector,
                keywords=case_keywords,
                top_k=top_k,
                keyword_weight=keyword_weight,
                vector_weight=vector_weight,
            )
            context = truncate(build_context(law_results), 2000)
            case_context = truncate(build_context(case_results), 2000)
            yield emit_progress(30)
            kb_brief = [
                {
                    "title": r.get("article_number") or "法条",
                    "snippet": (r.get("content") or "")[:140],
                    "detail": (r.get("content") or "")[:600],
                }
                for r in law_results[:3]
            ]
            case_brief = [
                {
                    "title": f"【案例】{r.get('title') or r.get('article_number') or '案例'}",
                    "snippet": r.get("summary") or (r.get("content") or "")[:240],
                    "detail": (r.get("content") or "")[:1200],
                }
                for r in case_results[:3]
            ]
            kb_brief = kb_brief + case_brief
            doc_pack = generate_document_with_model(
                api_key=api_key,
                base_url=base_url,
                model=model,
                user_query=user_query,
                memory_context=memory_context,
                file_text=file_text,
                file_assets=file_assets,
                doc_type=plan_doc_type or "通用文书",
                kb_context=(context + "\n\n" + case_context).strip(),
            )
            final_answer = doc_pack.get("answer") or "已生成文书草稿。"
            documents = doc_pack.get("documents") or []
            yield emit_progress(90)
            downloads = []
            for doc in documents:
                title = doc.get("title") or (plan_doc_type or "文书")
                body = doc.get("body") or ""
                images = doc.get("images") or []
                if not images and file_assets:
                    images = file_assets
                url = generate_pdf(title, body, images=images, doc_type=plan_doc_type or "文书")
                downloads.append({"title": title, "url": url})
            for i in range(0, len(final_answer), 24):
                chunk = final_answer[i : i + 24]
                yield json.dumps({"type": "delta", "text": chunk}, ensure_ascii=False) + "\n"
                time.sleep(0.02)
            doc_titles = [d.get("title") or "文书" for d in documents]
            for entry in extract_sparse_memory(user_query, final_answer, doc_titles=doc_titles):
                append_memory(session_id, entry)
            plan_detail = annotate_plan_detail(
                plan_detail,
                rag_keywords=[],
                search_results=[],
                documents=documents,
                action="generate_document",
                kb_brief=kb_brief,
                web_brief=[],
            )
            yield json.dumps({"type": "mindmap_loading"}, ensure_ascii=False) + "\n"
            case_analyses: List[Dict[str, Any]] = []
            try:
                cur = progress_val
                with ThreadPoolExecutor(max_workers=2) as ex:
                    fut = ex.submit(
                        generate_mindmap,
                        api_key=api_key,
                        base_url=base_url,
                        model=model,
                        user_query=user_query,
                        final_answer=final_answer,
                        plan_detail=plan_detail,
                        kb_brief=kb_brief,
                        case_brief=case_brief,
                        web_brief=[],
                        action="generate_document",
                    )
                    fut_case = None
                    if case_results:
                        fut_case = ex.submit(
                            batch_analyze_cases,
                            api_key=api_key,
                            base_url=base_url,
                            model=model,
                            user_query=user_query,
                            cases=case_results,
                            max_cases=min(3, len(case_results)),
                        )
                    while not fut.done():
                        cur = min(99, cur + 1)
                        yield emit_progress(cur)
                        time.sleep(0.6)
                    mindmap = fut.result()
                    if fut_case is not None:
                        try:
                            case_analyses = fut_case.result()
                        except Exception as ce:
                            logger.warning("Case analysis precompute failed: %s", ce)
            except Exception as e:
                logger.warning("Mindmap generation failed: %s", e)
                mindmap = {}
            if not mindmap:
                mindmap = build_fallback_mindmap(
                    plan_detail,
                    kb_brief=kb_brief,
                    case_brief=case_brief,
                    web_brief=[],
                    action="generate_document",
                    user_query=user_query,
                )
            yield emit_progress(100)
            meta = {
                "type": "meta",
                "keywords": rag_keywords,
                "results": law_results,
                "law_results": law_results,
                "case_results": case_results,
                "search_results": [],
                "highlight_keywords": [],
                "plan_detail": plan_detail,
                "downloads": downloads,
                "documents": documents,
                "mindmap": mindmap,
                "case_analyses": case_analyses,
            }
            yield json.dumps(meta, ensure_ascii=False) + "\n"
            return

        # 3) RAG search
        yield gen_status("模型检索：知识库搜索")
        kw_pack = generate_search_keywords(
            api_key=api_key,
            base_url=base_url,
            model=model,
            user_query=user_query,
            memory_context=memory_context,
        )
        law_keywords = kw_pack.get("law_keywords", [])
        case_keywords = kw_pack.get("case_keywords", [])
        rag_keywords = list(dict.fromkeys(law_keywords + case_keywords))
        query_vector = kw_pack.get("vector_query", user_query)
        web_queries = kw_pack.get("web_queries", [])
        law_results = retrieve_docs(
            DOCS,
            VECTORIZER,
            MATRIX,
            query_vector=query_vector,
            keywords=law_keywords,
            top_k=top_k,
            keyword_weight=keyword_weight,
            vector_weight=vector_weight,
        )
        case_results = retrieve_docs(
            CASE_DOCS,
            CASE_VECTORIZER,
            CASE_MATRIX,
            query_vector=query_vector,
            keywords=case_keywords,
            top_k=top_k,
            keyword_weight=keyword_weight,
            vector_weight=vector_weight,
        )
        context = truncate(build_context(law_results), CONFIG.max_context_chars)
        case_context = truncate(build_context(case_results), CONFIG.max_context_chars)
        yield emit_progress(30)

        # 4) decide web search based on plan
        search_needed = bool(plan.get("search_needed", False))
        if force_web_search:
            search_needed = True
        if action in ("rag", "direct_answer"):
            search_needed = False
        # heuristic: if user explicitly asks to查阅/规定/官网/政策/最新
        if re.search(r"查阅|规定|官网|政策|最新|标准|条例", user_query):
            search_needed = True
        if not search_needed:
            yield emit_progress(40)

        search_results = []
        queries = web_queries or []
        plan_queries = plan.get("search_queries", []) or []
        for q in plan_queries:
            if q and q not in queries:
                queries.append(q)
        search_target = int(payload.get("search_target") or 20)
        if search_target < 1:
            search_target = 1
        progress_count = 0
        if search_needed:
            yield emit_progress(35)
        max_rounds = 3
        round_idx = 0
        while search_needed and round_idx < max_rounds:
            round_idx += 1
            if not queries:
                queries = refine_search_queries(api_key, base_url, model, user_query, queries)
            if not queries:
                queries = heuristic_queries(user_query)
            if not queries:
                queries = [user_query]
            yield gen_status(f"模型搜索：网页检索中（第{round_idx}轮）")
            yield gen_trace("搜索关键词", ", ".join(queries) or "无")
            search_results = []
            for item in web_search_iter(queries, max_results=search_target, fetch_text=True, region=site_region):
                search_results.append(item)
                progress_count += 1
                pct = int(min(70, 35 + (progress_count / search_target) * 35))
                yield emit_progress(pct)
            if site_region == "cn_priority" and not search_results:
                for item in web_search_iter(queries, max_results=search_target, fetch_text=True, region="global"):
                    search_results.append(item)
                    progress_count += 1
                    pct = int(min(70, 35 + (progress_count / search_target) * 35))
                    yield emit_progress(pct)
            # deep fetch for top 1-2 results
            for i in range(min(2, len(search_results))):
                url = search_results[i].get("url", "")
                if url:
                    search_results[i]["content"] = fetch_page_text_deep(url)
            if search_results and evaluate_search_sufficiency(api_key, base_url, reasoning_model, user_query, search_results):
                break
            yield gen_trace("搜索结果", "本轮不足，正在调整关键词")
            queries = refine_search_queries(api_key, base_url, model, user_query, queries)
            if not queries:
                queries = heuristic_queries(user_query)

        if search_results:
            yield gen_status("模型搜索：结果分析")
            yield gen_trace("网页结果数", str(len(search_results)))
        else:
            yield gen_trace("网页搜索", "未找到结果，将基于知识库回答")

        # 4) execute with chat model
        yield gen_status("模型总结：生成最终答复")
        answer_start = 70 if search_needed else 40
        system_prompt = (
            "你是执行型智能体。根据计划、检索结果和用户问题输出最终方案。"
            "请在答案中引用相关法条与相似案例（若有）。"
            "只输出JSON。"
        )
        memory_context = build_memory_context(session_id)
        kb_brief = [
            {
                "title": r.get("article_number") or "法条",
                "snippet": (r.get("content") or "")[:140],
                "detail": (r.get("content") or "")[:600],
            }
            for r in law_results[:3]
        ]
        case_brief = [
            {
                "title": f"【案例】{r.get('title') or r.get('article_number') or '案例'}",
                "snippet": r.get("summary") or (r.get("content") or "")[:240],
                "detail": (r.get("content") or "")[:1200],
            }
            for r in case_results[:3]
        ]
        kb_brief = kb_brief + case_brief
        case_index_list = [
            f"{i+1}. {r.get('title') or r.get('article_number') or '案例'}"
            for i, r in enumerate(case_results[:5])
        ]
        web_index_list = [
            f"{i+1}. {(r.get('title') or '网页')} | {r.get('url') or ''}".strip()
            for i, r in enumerate(search_results[:5])
        ]
        web_brief = [
            {
                "title": r.get("title") or "网页",
                "url": r.get("url") or "",
                "snippet": (r.get("content") or r.get("snippet") or "")[:160],
                "detail": (r.get("content") or r.get("snippet") or "")[:800],
            }
            for r in search_results[:3]
        ]
        user_prompt = (
            f"用户问题：{user_query}\n\n"
            f"探索计划：{plan.get('plan', [])}\n\n"
            f"RAG法条：\n{context}\n\n"
            f"RAG案例：\n{case_context}\n\n"
            f"案例编号（引用时请使用【案例1】格式）：\n" + ("\n".join(case_index_list) or "无") + "\n\n"
            f"网页编号（引用时请使用【网页1】格式）：\n" + ("\n".join(web_index_list) or "无") + "\n\n"
            f"网页搜索结果：\n{json.dumps(search_results, ensure_ascii=False)}\n\n"
            f"用户补充材料：{file_text}\n\n"
            f"历史要点：{memory_context or '(无)'}\n\n"
            f"已上传图片资产ID：{file_assets}\n"
            "若需要在文书中使用图片，请在 documents[].images 中填写 asset_id，或在正文中插入 [[IMAGE:asset_id]]。\n\n"
            "注意：不要编造网页内容。若网页为空，避免冗长免责声明，直接给出可执行建议。\n"
            "请输出JSON：\n"
            "{\n"
            '  "answer": "最终答复，包含完整计划与行动建议",\n'
            '  "kb_entries": [{"content":"...","source":"...","keywords":["..."]}],\n'
            '  "documents": [{"title":"文书标题","body":"文书正文","images":["asset_id1","asset_id2"]}]\n'
            "}\n"
        )
        yield emit_progress(answer_start)
        try:
            cur = progress_val
            with ThreadPoolExecutor(max_workers=1) as ex:
                fut = ex.submit(
                    call_deepseek,
                    api_key=api_key,
                    base_url=base_url,
                    model=model,
                    messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
                    temperature=CONFIG.answer_temperature,
                    max_tokens=CONFIG.answer_max_tokens,
                )
                while not fut.done():
                    cur = min(89, cur + 1)
                    yield emit_progress(cur)
                    time.sleep(0.6)
                exec_text = fut.result()
        except Exception as e:
            raise e
        yield emit_progress(90)
        normalized = normalize_model_json(exec_text)
        exec_data = extract_json_block(normalized) or safe_json_loads(normalized) or {}
        if not isinstance(exec_data, dict):
            exec_data = {}
        final_answer = exec_data.get("answer") or extract_answer_field(normalized) or ""
        if not final_answer:
            final_answer = exec_text if not normalized.strip().startswith("{") else "未能解析结构化答复，请重试。"
        final_answer = sanitize_answer_text(final_answer)
        kb_entries = exec_data.get("kb_entries") or []
        documents = exec_data.get("documents") or []

        downloads = []
        if documents:
            for doc in documents:
                title = doc.get("title") or "文书"
                body = doc.get("body") or ""
                images = doc.get("images") or []
                if not images and file_assets:
                    images = file_assets
                url = generate_pdf(title, body, images=images, doc_type=title)
                downloads.append({"title": title, "url": url})

        if plan.get("should_update_kb") and kb_entries:
            update_extension_kb(kb_entries)

        for i in range(0, len(final_answer), 24):
            chunk = final_answer[i : i + 24]
            yield json.dumps({"type": "delta", "text": chunk}, ensure_ascii=False) + "\n"
            time.sleep(0.02)
        yield emit_progress(90)
        doc_titles = [d.get("title") or "文书" for d in documents] if documents else []
        for entry in extract_sparse_memory(user_query, final_answer, doc_titles=doc_titles):
            append_memory(session_id, entry)
        plan_detail = annotate_plan_detail(
            plan_detail,
            rag_keywords=rag_keywords,
            search_results=search_results,
            documents=documents,
            action=action,
            kb_brief=kb_brief,
            web_brief=web_brief,
        )
        yield json.dumps({"type": "mindmap_loading"}, ensure_ascii=False) + "\n"
        case_analyses: List[Dict[str, Any]] = []
        try:
            cur = progress_val
            with ThreadPoolExecutor(max_workers=2) as ex:
                fut = ex.submit(
                    generate_mindmap,
                    api_key=api_key,
                    base_url=base_url,
                    model=model,
                    user_query=user_query,
                    final_answer=final_answer,
                    plan_detail=plan_detail,
                    kb_brief=kb_brief,
                    case_brief=case_brief,
                    web_brief=web_brief,
                    action=action,
                )
                fut_case = None
                if case_results:
                    fut_case = ex.submit(
                        batch_analyze_cases,
                        api_key=api_key,
                        base_url=base_url,
                        model=model,
                        user_query=user_query,
                        cases=case_results,
                        max_cases=min(3, len(case_results)),
                    )
                while not fut.done():
                    cur = min(99, cur + 1)
                    yield emit_progress(cur)
                    time.sleep(0.6)
                mindmap = fut.result()
                if fut_case is not None:
                    try:
                        case_analyses = fut_case.result()
                    except Exception as ce:
                        logger.warning("Case analysis precompute failed: %s", ce)
        except Exception as e:
            logger.warning("Mindmap generation failed: %s", e)
            mindmap = {}
        if not mindmap:
            mindmap = build_fallback_mindmap(
                plan_detail,
                kb_brief=kb_brief,
                case_brief=case_brief,
                web_brief=web_brief,
                action=action,
                user_query=user_query,
            )
        yield emit_progress(100)
        meta = {
            "type": "meta",
            "keywords": rag_keywords,
            "results": law_results,
            "law_results": law_results,
            "case_results": case_results,
            "search_results": search_results,
            "highlight_keywords": web_queries or queries or [],
            "plan_detail": plan_detail,
            "downloads": downloads,
            "documents": documents,
            "mindmap": mindmap,
            "case_analyses": case_analyses,
        }
        yield json.dumps(meta, ensure_ascii=False) + "\n"

    return StreamingResponse(gen(), media_type="text/plain")
